#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word (.docx/.doc) → PDF for IPCS SDD delivery.

Windows (recommended): Microsoft Word COM export — preserves CJK table text and
complex OOXML (HTML function tables, embedded SVG) that LibreOffice often corrupts
into digits-only or blank cells.

Fallback: LibreOffice headless with a sanitized temp copy (table cells forced to
Normal style, numPr stripped).

Usage:
    python word_to_pdf.py --Word final_sdd.docx -o final.pdf
    python word_to_pdf.py --Word final_sdd.docx -o final.pdf --backend word
"""

from __future__ import annotations

import argparse
import logging
import re
import shutil
import subprocess
import sys
import tempfile
import time
from pathlib import Path
from typing import Optional

LOG_FORMAT = "%(asctime)s - [%(levelname)s] - %(message)s"
DEFAULT_TIMEOUT_SEC = 300
PDF_MARKERS = ("\u51fd\u6570\u539f\u578b", "Function prototype", "SWU_IPCS")
MIN_CJK_CHARS = 500

logging.basicConfig(level=logging.INFO, format=LOG_FORMAT)
logger = logging.getLogger("PDFConverter")


class ConversionError(Exception):
    pass


class DependencyError(ConversionError):
    pass


def resolve_output_path(input_file: Path, output_path: Optional[str]) -> Path:
    default_name = input_file.stem + ".pdf"
    if not output_path:
        return Path.cwd() / default_name
    dest = Path(output_path).resolve()
    if dest.is_dir():
        return dest / default_name
    if not dest.parent.exists():
        dest.parent.mkdir(parents=True, exist_ok=True)
    return dest


def validate_pdf_content(pdf_path: Path) -> None:
    """Fail fast if PDF looks like a broken LO export (no CJK / no function tables)."""
    text = ""
    try:
        import fitz  # pymupdf

        doc = fitz.open(str(pdf_path))
        for page in doc:
            text += page.get_text()
        doc.close()
    except ImportError:
        try:
            import pdfplumber

            with pdfplumber.open(str(pdf_path)) as pdf:
                for page in pdf.pages[: min(80, len(pdf.pages))]:
                    text += page.extract_text() or ""
        except ImportError:
            logger.warning("Skip PDF content validation (install pymupdf or pdfplumber).")
            return

    cjk = len(re.findall(r"[\u4e00-\u9fff]", text))
    if cjk < MIN_CJK_CHARS:
        raise ConversionError(
            f"PDF validation failed: only {cjk} CJK chars (need >={MIN_CJK_CHARS}). "
            "Table text likely missing — retry with --backend word."
        )
    if not any(m in text for m in PDF_MARKERS):
        raise ConversionError(
            "PDF validation failed: expected function-table markers not found in text."
        )
    logger.info("PDF content validation PASS (CJK chars=%d)", cjk)


def prepare_docx_for_libreoffice(src: Path, dst: Path) -> None:
    """Sanitize table paragraphs so LibreOffice does not render list numbers only."""
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(str(src))
    try:
        normal_id = doc.styles["Normal"].style_id
    except KeyError:
        normal_id = "Normal"

    fixed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    p_pr = para._element.get_or_add_pPr()
                    for num_pr in p_pr.findall(qn("w:numPr")):
                        p_pr.remove(num_pr)
                        fixed += 1
                    p_style = p_pr.find(qn("w:pStyle"))
                    if p_style is None:
                        p_style = OxmlElement("w:pStyle")
                        p_pr.insert(0, p_style)
                    if p_style.get(qn("w:val")) != normal_id:
                        p_style.set(qn("w:val"), normal_id)
                        fixed += 1
    doc.save(str(dst))
    logger.info("Prepared LibreOffice temp docx (table paragraph fixes=%d)", fixed)


class WordComConverter:
    """Export via installed Microsoft Word (Windows)."""

    def is_available(self) -> bool:
        if sys.platform != "win32":
            return False
        try:
            import win32com.client  # noqa: F401

            return True
        except ImportError:
            return False

    def convert(self, input_file: Path, output_file: Path) -> None:
        import win32com.client

        logger.info("Converting with Microsoft Word COM: %s", input_file.name)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        start = time.time()
        try:
            doc = word.Documents.Open(str(input_file.resolve()))
            try:
                if output_file.exists():
                    logger.info("Overwriting existing file: %s", output_file)
                # wdExportFormatPDF = 17, wdExportCreateHeadingBookmarks = 1
                doc.ExportAsFixedFormat(
                    OutputFileName=str(output_file.resolve()),
                    ExportFormat=17,
                    OpenAfterExport=False,
                    OptimizeFor=0,
                    CreateBookmarks=1,
                    DocStructureTags=True,
                    BitmapMissingFonts=True,
                )
            finally:
                doc.Close(False)
        finally:
            word.Quit()
        if not output_file.exists():
            raise ConversionError("Word COM finished but PDF was not created.")
        logger.info("Word COM conversion finished in %.2fs", time.time() - start)


class LibreOfficeConverter:
    def __init__(self, libreoffice_path: Optional[str] = None):
        self.libreoffice_exec = libreoffice_path or self._detect_libreoffice()
        if not self.libreoffice_exec:
            raise DependencyError("LibreOffice executable not found.")

    @staticmethod
    def _detect_libreoffice() -> Optional[str]:
        candidates = [
            "libreoffice",
            "soffice",
            "libreoffice7.6",
            "libreoffice7.5",
            "/usr/bin/libreoffice",
            "/usr/bin/soffice",
        ]
        if sys.platform == "win32":
            candidates.extend(
                [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                ]
            )
        for cmd in candidates:
            if shutil.which(cmd) or Path(cmd).exists():
                return cmd
        return None

    def convert(self, input_file: Path, output_file: Path) -> None:
        with tempfile.TemporaryDirectory(prefix="lo_pdf_src_") as tmp_src:
            prepared = Path(tmp_src) / input_file.name
            prepare_docx_for_libreoffice(input_file, prepared)
            with tempfile.TemporaryDirectory(prefix="lo_convert_") as user_dir:
                with tempfile.TemporaryDirectory(prefix="lo_out_") as out_dir:
                    user_uri = Path(user_dir).as_uri()
                    filter_opts = (
                        'pdf:writer_pdf_Export:'
                        '{"EmbedStandardFonts":{"type":"boolean","value":"true"},'
                        '"UseLosslessCompression":{"type":"boolean","value":"true"}}'
                    )
                    cmd = [
                        self.libreoffice_exec,
                        f"-env:UserInstallation={user_uri}",
                        "--headless",
                        "--convert-to",
                        filter_opts,
                        "--outdir",
                        out_dir,
                        str(prepared),
                    ]
                    logger.info("Converting with LibreOffice: %s", input_file.name)
                    start = time.time()
                    try:
                        subprocess.run(
                            cmd,
                            stdout=subprocess.PIPE,
                            stderr=subprocess.PIPE,
                            timeout=DEFAULT_TIMEOUT_SEC,
                            check=True,
                        )
                    except subprocess.TimeoutExpired as exc:
                        raise ConversionError("LibreOffice conversion timed out.") from exc
                    except subprocess.CalledProcessError as exc:
                        err = (exc.stderr or exc.stdout or b"").decode(
                            "utf-8", errors="replace"
                        )
                        raise ConversionError(
                            f"LibreOffice failed (exit {exc.returncode}): {err.strip()}"
                        ) from exc
                    logger.info(
                        "LibreOffice conversion finished in %.2fs", time.time() - start
                    )
                    temp_pdf = Path(out_dir) / (prepared.stem + ".pdf")
                    if not temp_pdf.exists():
                        raise ConversionError("LibreOffice did not produce a PDF.")
                    if output_file.exists():
                        logger.info("Overwriting existing file: %s", output_file)
                    shutil.move(str(temp_pdf), str(output_file))


def convert_docx_to_pdf(
    input_path: str | Path,
    output_path: Optional[str | Path] = None,
    *,
    backend: str = "auto",
    validate: bool = True,
) -> Path:
    input_file = Path(input_path).resolve()
    if not input_file.is_file():
        raise FileNotFoundError(f"Input file not found: {input_file}")
    output_file = resolve_output_path(input_file, str(output_path) if output_path else None)

    word = WordComConverter()
    if backend in ("auto", "word"):
        if word.is_available():
            word.convert(input_file, output_file)
            if validate:
                validate_pdf_content(output_file)
            logger.info("Success! Output saved to: %s", output_file)
            return output_file
        if backend == "word":
            raise DependencyError(
                "Microsoft Word COM unavailable (need Windows + pywin32 + Word)."
            )
        logger.warning("Word COM unavailable; falling back to LibreOffice.")

    lo = LibreOfficeConverter()
    lo.convert(input_file, output_file)
    if validate:
        validate_pdf_content(output_file)
    logger.info("Success! Output saved to: %s", output_file)
    return output_file


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert Word to PDF (Word COM on Windows, LibreOffice fallback)."
    )
    parser.add_argument("--Word", dest="input_file", required=True, help="Input .docx/.doc")
    parser.add_argument("-o", "--output", dest="output_file", help="Output .pdf path")
    parser.add_argument(
        "--backend",
        choices=("auto", "word", "libreoffice"),
        default="auto",
        help="Conversion backend (default: auto → Word on Windows)",
    )
    parser.add_argument(
        "--no-validate",
        action="store_true",
        help="Skip post-export PDF text validation",
    )
    args = parser.parse_args()
    try:
        convert_docx_to_pdf(
            args.input_file,
            args.output_file,
            backend=args.backend,
            validate=not args.no_validate,
        )
    except (ConversionError, DependencyError, FileNotFoundError) as exc:
        logger.error("Conversion Error: %s", exc)
        sys.exit(1)
    except Exception:
        logger.exception("Unexpected Error")
        sys.exit(1)


if __name__ == "__main__":
    main()
