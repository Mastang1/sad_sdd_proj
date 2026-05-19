# -*- coding: utf-8 -*-
r"""
将 ``md_sdd_0519.md`` 中 **OSAL(THREADX) 函数实现** 章节同步到 TF（``final_sdd.docx``）。

依照工作区 ``.cursorrules`` / ``prompt_docx-to-md.md`` task-3：

1. 在 TF 中定位 ``Heading 2`` 标题含 ``OSAL(THREADX)`` 的小节，删除该标题与下一 ``Heading 2``
   （``Dynamic Detailed Design``）之间的全部表格、``processing flow`` 段与流程图插图，**不改动**其他章节。
2. 自 Markdown 提取各函数的 HTML 表格与 ``flow_svgs/tx_3_4_*.svg`` 路径，按源码顺序插入：
   函数名（``Heading 3``）→ Word 表格（合并/拆分与 MD 一致）→ ``processing flow`` 正文段 → SVG 转 PNG 插图（宽 14 cm、居中）。
3. 更新 THREADX 小节标题，去掉 ``（todo）`` 字样。
4. 对新插入表格/插图应用 ``format_final_sdd`` 中的边框、单元格 10 pt 字体与插图版式规则。

**依赖**::

    pip install python-docx cairosvg beautifulsoup4 lxml

**使用方式**（在仓库根目录 ``c1_sdd`` 下）::

    python format_docx_py/sync_osal_threadx_to_tf.py
    python format_docx_py/sync_osal_threadx_to_tf.py "D:\\path\\final_sdd.docx"

若 ``final_sdd.docx`` 被 Word 占用，将写入 ``final_sdd_threadx_sync.docx`` 并提示手动替换。
"""

from __future__ import annotations

import io
import re
import sys
from pathlib import Path

import cairosvg
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

from format_final_sdd import (
    DEFAULT_TF,
    _BODY_FONT_PT,
    _DEFAULT_FIGURE_WIDTH_CM,
    _EAST_ASIA_FONT,
    _TABLE_FONT_PT,
    _WESTERN_FONT,
    apply_solid_table_borders,
    format_all_table_cells,
    format_body_paragraphs,
    format_figure_chart_paragraphs_layout,
    scale_inline_shapes_margin_and_width,
    _set_run_font,
)

WORKSPACE = Path(__file__).resolve().parents[1]
DEFAULT_MD = WORKSPACE / "md_sdd_0519.md"
SECTION_H2 = "OSAL(THREADX)"
SECTION_END_H2 = "Dynamic Detailed Design"
H2_NEW_TITLE = "OSAL(THREADX) 函数实现"
FIGURE_WIDTH_CM = _DEFAULT_FIGURE_WIDTH_CM

_RE_FUNC_BLOCK = re.compile(
    r"### 3\.4\.\d+ (\w+)\s*\n+<table.*?</table>\s*\n+processing flow\s*\n+!\[[^\]]*\]\(([^)]+)\)",
    re.DOTALL,
)


def _tf_path() -> Path:
    if len(sys.argv) > 1:
        return Path(sys.argv[1]).resolve()
    return DEFAULT_TF.resolve()


def _plain_text(p_el) -> str:
    return "".join(t.text or "" for t in p_el.iter(qn("w:t"))).strip()


def _paragraph_style_name(doc: Document, p_el) -> str | None:
    ppr = p_el.pPr
    if ppr is None or ppr.pStyle is None:
        return None
    try:
        from docx.enum.style import WD_STYLE_TYPE

        return doc.styles.get_by_id(ppr.pStyle.val, WD_STYLE_TYPE.PARAGRAPH).name
    except (KeyError, ValueError, AttributeError):
        return None


def _is_h2(doc: Document, p_el, needle: str) -> bool:
    return _paragraph_style_name(doc, p_el) == "Heading 2" and needle in _plain_text(p_el)


def parse_html_placements(table_html: str):
    soup = BeautifulSoup(table_html, "html.parser")
    table = soup.find("table")
    if not table:
        return 0, 0, []
    occupied: dict[tuple[int, int], bool] = {}
    placements: list[tuple[int, int, int, int, str]] = []
    cur_r = 0
    for tr in table.find_all("tr"):
        tds = tr.find_all("td")
        if not tds:
            continue
        cur_c = 0
        for td in tds:
            while (cur_r, cur_c) in occupied:
                cur_c += 1
            rs = int(td.get("rowspan", 1))
            cs = int(td.get("colspan", 1))
            text = td.get_text(separator="\n").strip()
            placements.append((cur_r, cur_c, rs, cs, text))
            for dr in range(rs):
                for dc in range(cs):
                    occupied[(cur_r + dr, cur_c + dc)] = True
            cur_c += cs
        cur_r += 1
    nrows = cur_r
    ncols = max((c + cs for _, c, _, cs, _ in placements), default=0)
    return nrows, ncols, placements


def load_threadx_blocks(md_path: Path) -> list[tuple[str, str, Path]]:
    text = md_path.read_text(encoding="utf-8")
    start = text.find("## OSAL(THREADX)")
    if start < 0:
        raise SystemExit("md 中未找到 ## OSAL(THREADX)")
    end = text.find("## Dynamic Detailed Design", start)
    if end < 0:
        raise SystemExit("md 中未找到 ## Dynamic Detailed Design")
    chunk = text[start:end]
    blocks: list[tuple[str, str, Path]] = []
    for fn, img_rel in _RE_FUNC_BLOCK.findall(chunk):
        svg = (WORKSPACE / img_rel.replace("/", "\\")).resolve()
        if not svg.is_file():
            raise SystemExit(f"缺少 SVG：{svg}")
        m = re.search(
            rf"### 3\.4\.\d+ {re.escape(fn)}\s*\n+(<table.*?</table>)",
            chunk,
            re.DOTALL,
        )
        if not m:
            raise SystemExit(f"未找到 {fn} 的表格")
        blocks.append((fn, m.group(1), svg))
    if len(blocks) != 7:
        raise SystemExit(f"期望 7 个函数块，实际 {len(blocks)}：{[b[0] for b in blocks]}")
    return blocks


def svg_to_png(svg_path: Path) -> bytes:
    return cairosvg.svg2png(url=str(svg_path), dpi=120)


def insert_table_after(doc: Document, anchor_el, nrows: int, ncols: int, placements):
    temp = doc.add_table(rows=nrows, cols=ncols)
    tbl_el = temp._tbl
    temp._tbl.getparent().remove(tbl_el)
    anchor_el.addnext(tbl_el)
    table = Table(tbl_el, doc)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r, c, rs, cs, text in placements:
        cell = table.cell(r, c)
        if rs > 1 or cs > 1:
            cell.merge(table.cell(r + rs - 1, c + cs - 1))
        cell.text = text or " "
        for p in cell.paragraphs:
            for run in p.runs:
                _set_run_font(
                    run,
                    western=_WESTERN_FONT,
                    east_asia=_EAST_ASIA_FONT,
                    size_pt=_TABLE_FONT_PT,
                    bold=False,
                )
    apply_solid_table_borders(table)
    return tbl_el


def _apply_body11(para: Paragraph) -> None:
    pf = para.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = None


def insert_paragraph_after(
    doc: Document, anchor_el, text: str, *, style: str | None = None, center: bool = False
):
    p = doc.add_paragraph()
    if style:
        try:
            p.style = doc.styles[style]
        except KeyError:
            pass
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        run = p.add_run(text)
        _set_run_font(
            run,
            western=_WESTERN_FONT,
            east_asia=_EAST_ASIA_FONT,
            size_pt=int(_BODY_FONT_PT),
            bold=False,
        )
    _apply_body11(p)
    el = p._element
    el.getparent().remove(el)
    anchor_el.addnext(el)
    return el


def insert_figure_after(doc: Document, anchor_el, png_bytes: bytes):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(io.BytesIO(png_bytes), width=Cm(FIGURE_WIDTH_CM))
    _apply_body11(p)
    el = p._element
    el.getparent().remove(el)
    anchor_el.addnext(el)
    return el


def find_threadx_region(doc: Document, body):
    h2_start = None
    h2_end = None
    for ch in body:
        if ch.tag != qn("w:p"):
            continue
        if _is_h2(doc, ch, SECTION_H2):
            h2_start = ch
        elif h2_start is not None and _is_h2(doc, ch, SECTION_END_H2):
            h2_end = ch
            break
    if h2_start is None or h2_end is None:
        raise SystemExit("TF 中未定位 OSAL(THREADX) 或 Dynamic Detailed Design 标题")
    return h2_start, h2_end


def clear_between(body, start_el, end_el) -> int:
    parent = start_el.getparent()
    kids = list(parent)
    i0 = kids.index(start_el) + 1
    i1 = kids.index(end_el)
    removed = 0
    for el in kids[i0:i1]:
        parent.remove(el)
        removed += 1
    return removed


def set_h2_title(doc: Document, h2_el) -> None:
    """将 THREADX 小节 H2 全文替换为固定标题（避免多 run 残留）。"""
    para = Paragraph(h2_el, doc._body)
    para.text = H2_NEW_TITLE


def sync_section(doc: Document, blocks: list[tuple[str, str, Path]]) -> None:
    body = doc.element.body
    h2_start, h2_end = find_threadx_region(doc, body)
    n_rm = clear_between(body, h2_start, h2_end)
    set_h2_title(doc, h2_start)
    anchor = h2_start
    for fn, table_html, svg_path in blocks:
        anchor = insert_paragraph_after(doc, anchor, fn, style="Heading 3")
        nrows, ncols, placements = parse_html_placements(table_html)
        if nrows == 0 or ncols == 0:
            raise SystemExit(f"{fn} 表格解析失败")
        anchor = insert_table_after(doc, anchor, nrows, ncols, placements)
        anchor = insert_paragraph_after(doc, anchor, "processing flow", style="Normal")
        png = svg_to_png(svg_path)
        anchor = insert_figure_after(doc, anchor, png)
    print(f"  删除旧块元素: {n_rm}")
    print(f"  插入函数: {len(blocks)}")


def main() -> None:
    tf = _tf_path()
    md = DEFAULT_MD
    if not md.is_file():
        sys.exit(f"未找到 Markdown：{md}")
    if not tf.is_file():
        sys.exit(f"未找到 TF：{tf}")

    blocks = load_threadx_blocks(md)
    doc = Document(str(tf))
    sync_section(doc, blocks)
    format_all_table_cells(doc)
    format_body_paragraphs(doc)
    format_figure_chart_paragraphs_layout(doc)
    scale_inline_shapes_margin_and_width(doc)

    out = tf
    try:
        doc.save(str(tf))
    except PermissionError:
        out = tf.with_name("final_sdd_threadx_sync.docx")
        doc.save(str(out))
        print(f"NOTE: {tf.name} 被占用，已保存为 {out.name}")

    print("同步完成：", out)
    for fn, _, svg in blocks:
        print(f"  - {fn} <- {svg.name}")


if __name__ == "__main__":
    main()
