# -*- coding: utf-8 -*-
"""
Build final_sdd.docx from ipcs_sdd.md:
- Preserve content in final_sdd.docx up to the first paragraph that contains a
  manual Page Break — remove everything after that page break paragraph;
  if no page break is found, insert one after the first body block so the merged
  body always follows a cover page boundary.
- Append pandoc-rendered Markdown (pipe + HTML tables, images).
- Render ```mermaid blocks to SVG: prefer local `@mermaid-js/mermaid-cli` (mmdc),
  otherwise https://mermaid.ink/svg/...
- Pandoc/docx embedding on Windows lacks rsvg-convert: rasterize Markdown ``*.svg``
  refs to `_docx_raster/**.png`` via CairoSVG (`pip install cairosvg`) before Pandoc.
"""
from __future__ import annotations

import base64
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import urllib.error
import urllib.request
import zlib
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer

ROOT = Path(__file__).resolve().parents[1]
MERMAID_INK_UA = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36"
)
MERMAID_INK_MAX_PLAIN_URL = 7900
SOURCE_MD = ROOT / "ipcs_sdd.md"
TARGET_DOCX = ROOT / "final_sdd.docx"
MERMAID_DIR = ROOT / "mermaid_svgs"
PANDOC_SRC = ROOT / "_pandoc_for_word.md"
BODY_DOCX = ROOT / "_body_generated.docx"

try:
    import cairosvg
except ImportError:
    cairosvg = None

DOCX_RASTER = ROOT / "_docx_raster"


def make_hard_page_break_paragraph():
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def first_paragraph_containing_manual_page_break(body) -> OxmlElement | None:
    for child in body:
        if child.tag != qn("w:p"):
            continue
        for br in child.iter(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return child
    return None


def trim_keeps_cover_removes_rest(doc: Document) -> None:
    body = doc.element.body
    p_break = first_paragraph_containing_manual_page_break(body)
    if p_break is None:
        kids = list(body)
        if not kids:
            doc.add_paragraph("（封面占位）")
            p = doc.add_paragraph()
            from docx.enum.text import WD_BREAK

            p.add_run().add_break(WD_BREAK.PAGE)
            return
        print(
            "[warn] 未发现手动分页符，已在第一个正文块之后自动插入分页符作为封面分界线。"
        )
        p_break = make_hard_page_break_paragraph()
        kids[0].addnext(p_break)
    nx = p_break.getnext()
    while nx is not None:
        nxt = nx.getnext()
        body.remove(nx)
        nx = nxt


def ensure_default_cover(docx_path: Path) -> None:
    if docx_path.is_file():
        return
    print(f"[info] Creating default cover template → {docx_path}")
    MERMAID_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("IPCS Driver 软件详细设计说明书", level=0)
    doc.add_paragraph("Detailed Design Specification (SWE.3)")
    doc.add_paragraph("")
    doc.add_paragraph(
        "（自动生成扉页占位。请替换为贵司正式封面，并在封面最后一栏之后插入 Word「分页符」。"
        "分页符后面的内容会在每次导出时被 Markdown 生成的正文替换。）"
    )
    p = doc.add_paragraph()
    run = p.add_run()
    from docx.enum.text import WD_BREAK

    run.add_break(WD_BREAK.PAGE)
    doc.save(docx_path)


def _mmdc_path() -> Path | None:
    bin_dir = ROOT / "node_modules" / ".bin"
    if sys.platform == "win32":
        cand = bin_dir / "mmdc.cmd"
    else:
        cand = bin_dir / "mmdc"
    if cand.is_file():
        return cand
    for key in ("mmdc", "mmdc.cmd"):
        w = shutil.which(key)
        if w:
            return Path(w)
    return None


def _render_mermaid_mmdc(mmd: Path, svg: Path, code: str) -> None:
    mmd.parent.mkdir(parents=True, exist_ok=True)
    mmd.write_text(code, encoding="utf-8")
    mmdc = _mmdc_path()
    if mmdc and mmdc.suffix.lower() == ".cmd":
        cmd = [
            os.environ.get("COMSPEC", "cmd.exe"),
            "/c",
            str(mmdc),
            "-i",
            str(mmd.resolve()),
            "-o",
            str(svg.resolve()),
            "-b",
            "transparent",
        ]
    elif mmdc:
        cmd = [
            str(mmdc),
            "-i",
            str(mmd.resolve()),
            "-o",
            str(svg.resolve()),
            "-b",
            "transparent",
        ]
    else:
        exe = shutil.which("npx") or shutil.which("npx.cmd")
        if not exe:
            raise FileNotFoundError("mmdc and npx not found")
        if exe.lower().endswith(".cmd") or sys.platform == "win32":
            cmd_str = (
                f'"{exe}" --yes @mermaid-js/mermaid-cli '
                f'-i "{mmd.resolve()}" -o "{svg.resolve()}" -b transparent'
            )
            cmd = [os.environ.get("COMSPEC", "cmd.exe"), "/c", cmd_str]
        else:
            cmd = [
                exe,
                "--yes",
                "@mermaid-js/mermaid-cli",
                "-i",
                str(mmd.resolve()),
                "-o",
                str(svg.resolve()),
                "-b",
                "transparent",
            ]
    subprocess.run(
        cmd,
        cwd=str(ROOT),
        check=True,
        stderr=subprocess.PIPE,
        stdout=subprocess.PIPE,
        timeout=180,
        text=True,
    )


def _render_mermaid_mermaid_ink(code: str, svg_path: Path) -> None:
    """SVG via https://mermaid.ink/svg/... (requires HTTP User-Agent).

    Uses a short plaintext base64 path when URL length allows,
    otherwise ``pako:<zlib(json)>`` compressed payload.

    Optionally set MERMAID_INK_BASE to a self-hosted base (scheme+host only),
    e.g. ``http://localhost:3000`` — requests become that host plus the path after
    ``https://mermaid.ink``.
    """
    base_hdr = {"User-Agent": MERMAID_INK_UA}
    raw_json = bytes(
        json.dumps({"code": code, "mermaid": {}}, separators=(",", ":")),
        encoding="utf-8",
    )
    pako_seg = (
        "pako:"
        + base64.urlsafe_b64encode(zlib.compress(raw_json, level=9)).decode("ascii").rstrip("=")
    )
    urls: list[tuple[str, str]] = [
        ("pako", f"https://mermaid.ink/svg/{pako_seg}")
    ]
    plain_seg = base64.urlsafe_b64encode(code.encode("utf-8")).decode("ascii").rstrip("=")
    u_plain = f"https://mermaid.ink/svg/{plain_seg}"
    if len(u_plain) <= MERMAID_INK_MAX_PLAIN_URL:
        urls.insert(0, ("plain", u_plain))

    proxy_base = os.environ.get("MERMAID_INK_BASE", "").rstrip("/")

    def rewrite_if_proxy(u: str) -> str:
        if not proxy_base or not u.startswith("https://mermaid.ink"):
            return u
        return proxy_base + u.removeprefix("https://mermaid.ink")

    last_err: BaseException | None = None

    for tag, candidate in urls:
        resolved = rewrite_if_proxy(candidate)
        req = urllib.request.Request(resolved, headers=base_hdr, method="GET")
        ok_url = False
        for attempt in range(5):
            try:
                with urllib.request.urlopen(req, timeout=180) as resp:
                    svg_path.write_bytes(resp.read())
                ok_url = True
                break
            except urllib.error.HTTPError as e:
                last_err = e
                payload = ""
                try:
                    payload = (e.fp.read() or b"").decode(errors="replace")[:240]
                except Exception:
                    pass
                print(
                    f"[info] mermaid.ink GET ({tag}) {e.code}: {payload!r}",
                    file=sys.stderr,
                )
                break
            except urllib.error.URLError as e:
                last_err = e
                wait = 0.4 * (attempt + 1)
                print(
                    f"[info] mermaid.ink ({tag}) attempt {attempt + 1}/5: {e}; "
                    f"sleeps {wait:.1f}s",
                    file=sys.stderr,
                )
                time.sleep(wait)
        if ok_url:
            return

    raise last_err or RuntimeError("mermaid.ink: no URLs attempted")


def preprocess_mermaid(md_text: str) -> str:
    MERMAID_DIR.mkdir(parents=True, exist_ok=True)
    pattern = re.compile(r"```mermaid\s*\n(.*?)```", re.DOTALL | re.IGNORECASE)
    ctr = {"n": 0, "ok": 0}

    def repl(m: re.Match[str]) -> str:
        ctr["n"] += 1
        n = ctr["n"]
        code = m.group(1).strip()
        stem = f"mermaid_{n:03d}"
        mmd = MERMAID_DIR / f"{stem}.mmd"
        svg = MERMAID_DIR / f"{stem}.svg"
        try:
            try:
                _render_mermaid_mmdc(mmd, svg, code)
            except (
                subprocess.CalledProcessError,
                subprocess.TimeoutExpired,
                OSError,
                FileNotFoundError,
            ) as e1:
                print(
                    f"[info] Local mermaid-cli unavailable or failed ({stem}): {e1}; "
                    "trying mermaid.ink..."
                )
                _render_mermaid_mermaid_ink(code, svg)
                mmd.write_text(code, encoding="utf-8")
        except (
            urllib.error.URLError,
            urllib.error.HTTPError,
            subprocess.CalledProcessError,
            subprocess.TimeoutExpired,
            OSError,
        ) as e:
            print(f"[error] Mermaid render failed ({stem}): {e}", file=sys.stderr)
            return f"[Mermaid rendering failed — see `{stem}.mmd`]\n```\n{code}\n```\n"

        ctr["ok"] += 1
        rel = str(svg.relative_to(ROOT)).replace("\\", "/")
        return f"![mermaid-{n}]({rel})\n"

    out = pattern.sub(repl, md_text)
    if ctr["n"]:
        print(
            f"[info] Rendered {ctr['ok']}/{ctr['n']} Mermaid diagram(s) → {MERMAID_DIR}/"
        )
    return out


def rasterize_svg_refs_for_docx(md_text: str, *, clear_cache: bool = True) -> str:
    """Replace local ``![alt](relative.svg)`` with ``_docx_raster/**/*.png``.

    Keeps authoritative SVGs on disk (`ipcs_sdd.md` still points at ``.svg`` in
    the working tree); Pandoc DOCX sees PNG so images embed without rsvg-convert.

    When ``clear_cache`` is False, existing raster files are kept so partial DOCX
    patches can regenerate only newly referenced diagrams.
    """

    if cairosvg is None:
        sys.exit(
            "pip install cairosvg — required to embed SVG-based diagrams into DOCX "
            "when rsvg-convert is not installed (typical on Windows)."
        )

    if clear_cache:
        shutil.rmtree(DOCX_RASTER, ignore_errors=True)
    DOCX_RASTER.mkdir(parents=True, exist_ok=True)
    scale = float(os.environ.get("DOCX_SVG_SCALE", "2"))
    md_img = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")

    def repl(m: re.Match[str]) -> str:
        alt = m.group(1)
        raw = m.group(2).strip().strip('"')
        if raw.startswith(("http://", "https://", "data:")):
            return m.group(0)
        rel_fs = Path(raw.replace("\\", "/"))
        posix = rel_fs.as_posix()
        if not posix.lower().endswith(".svg"):
            return m.group(0)
        svg_path = (ROOT / rel_fs).resolve()
        if not svg_path.is_file():
            print(f"[warn] rasterize: missing {svg_path}", file=sys.stderr)
            return m.group(0)
        out_rel_fs = Path("_docx_raster") / rel_fs.with_suffix(".png")
        out_path = (ROOT / out_rel_fs).resolve()
        out_path.parent.mkdir(parents=True, exist_ok=True)
        cairosvg.svg2png(url=str(svg_path), write_to=str(out_path), scale=scale)
        return f"![{alt}]({out_rel_fs.as_posix()})"

    return md_img.sub(repl, md_text)


def run_pandoc(md_path: Path, out_docx: Path) -> None:
    rpath = ";".join(
        [
            str(ROOT),
            str(ROOT / "flow_svgs"),
            str(MERMAID_DIR),
        ]
    )
    cmd = [
        "pandoc",
        str(md_path),
        "-o",
        str(out_docx),
        "-f",
        "markdown+pipe_tables+raw_html+smart+yaml_metadata_block",
        "-t",
        "docx",
        "--resource-path",
        rpath,
        "--standalone",
    ]
    subprocess.run(cmd, cwd=str(ROOT), check=True)
    print(f"[info] Pandoc → {out_docx}")


def _unlink_retry(path: Path, attempts: int = 12, delay_sec: float = 0.15) -> None:
    for _ in range(attempts):
        try:
            path.unlink(missing_ok=True)
            return
        except PermissionError:
            time.sleep(delay_sec)


def merge_cover_and_body(cover_path: Path, body_path: Path, out_path: Path) -> None:
    fd, cover_copy = tempfile.mkstemp(suffix=".docx", dir=str(ROOT))
    os.close(fd)
    cover_copy_p = Path(cover_copy)
    tmp = Path(tempfile.mkstemp(suffix=".docx", dir=str(ROOT))[1])
    try:
        shutil.copy2(cover_path, cover_copy_p)
        master = Document(str(cover_copy_p))
        trim_keeps_cover_removes_rest(master)
        composer = Composer(master)
        composer.append(Document(str(body_path)))
        composer.save(str(tmp))
        del composer, master
        try:
            shutil.copyfile(str(tmp), str(out_path))
            print(f"[info] Merged cover + body → {out_path}")
        except OSError as e:
            alt = out_path.with_name(out_path.stem + ".generated.docx")
            shutil.copyfile(str(tmp), str(alt))
            print(
                f"[warn] Cannot overwrite {out_path} ({e}); wrote {alt}. "
                "Close Word and copy/rename, or run again.",
                file=sys.stderr,
            )
    finally:
        _unlink_retry(tmp)
        _unlink_retry(cover_copy_p)


def main() -> None:
    if not SOURCE_MD.is_file():
        sys.exit(f"missing source: {SOURCE_MD}")
    ensure_default_cover(TARGET_DOCX)

    md = SOURCE_MD.read_text(encoding="utf-8")
    md = preprocess_mermaid(md)
    md = rasterize_svg_refs_for_docx(md, clear_cache=True)
    PANDOC_SRC.write_text(md, encoding="utf-8")

    run_pandoc(PANDOC_SRC, BODY_DOCX)
    merge_cover_and_body(TARGET_DOCX, BODY_DOCX, TARGET_DOCX)

    for p in (PANDOC_SRC, BODY_DOCX):
        p.unlink(missing_ok=True)

    print("[done]")


if __name__ == "__main__":
    main()
