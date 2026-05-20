# -*- coding: utf-8 -*-
r"""
MD → DOCX 一致性校验（task-6 第 4 项）。

对比 ``md_sdd_0519.md`` 与 ``final_sdd.docx`` 的结构与关键内容覆盖率，
输出 ``validate_md_docx_report.txt``。

使用方式（仓库根目录）::

    python cursor_tmp/format_docx_py/validate_md_docx_consistency.py
    python cursor_tmp/format_docx_py/validate_md_docx_consistency.py path/to/final_sdd.docx

退出码：0 = 无 ERROR 级问题；1 = 存在 ERROR。
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

_CURSOR_TMP = Path(__file__).resolve().parents[1]
_FORMAT_PY = _CURSOR_TMP / "format_docx_py"
for _p in (_CURSOR_TMP, _FORMAT_PY):
    if str(_p) not in sys.path:
        sys.path.insert(0, str(_p))
from workspace_paths import WORKSPACE_ROOT, MD_SDD_0519, FINAL_SDD_DOCX, VALIDATE_REPORT
from svg_extent_utils import parse_svg_viewbox_wh

ROOT = WORKSPACE_ROOT
MD_PATH = MD_SDD_0519
DOCX_PATH = FINAL_SDD_DOCX
REPORT_PATH = VALIDATE_REPORT

_HEADING_MD = re.compile(r"^(#{1,4})\s+(.+)$", re.MULTILINE)
_IMG_MD = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_TABLE_ROW = re.compile(r"^\|.+\|$", re.MULTILINE)
_TABLE_SEP = re.compile(r"^\|[\s\-:|]+\|$", re.MULTILINE)
_SWU = re.compile(r"SWU_IPCS_[A-Z0-9_]+")
_FUNC_ANCHOR = re.compile(r"^###\s+(\d+\.\d+\.\d+)\s+(\w+)", re.MULTILINE)


def _norm_heading(text: str) -> str:
    t = text.strip()
    t = re.sub(r"^\d+(\.\d+)*\s+", "", t)
    t = re.sub(r"\s+", " ", t)
    return t.lower()


def parse_md(md: str) -> dict:
    headings: list[tuple[int, str]] = []
    for m in _HEADING_MD.finditer(md):
        level = len(m.group(1))
        if m.group(1).startswith("#") and level <= 4:
            headings.append((level, _norm_heading(m.group(2))))

    images: list[str] = []
    missing_files: list[str] = []
    for _alt, path in _IMG_MD.findall(md):
        if path.startswith("http"):
            continue
        images.append(path)
        p = ROOT / path.replace("/", "\\") if "\\" not in path else ROOT / path
        if not p.is_file():
            missing_files.append(path)

    lines = md.splitlines()
    table_count = 0
    i = 0
    while i < len(lines):
        if _TABLE_ROW.match(lines[i]) and i + 1 < len(lines) and _TABLE_SEP.match(lines[i + 1]):
            table_count += 1
            i += 2
            while i < len(lines) and _TABLE_ROW.match(lines[i]):
                i += 1
            continue
        i += 1

    swu_ids = sorted(set(_SWU.findall(md)))
    func_sections = _FUNC_ANCHOR.findall(md)

    html_table_count = len(re.findall(r'<table border="1"', md))

    return {
        "headings": headings,
        "images": images,
        "missing_image_files": missing_files,
        "tables": table_count,
        "html_tables": html_table_count,
        "swu_ids": swu_ids,
        "func_sections": func_sections,
    }


def _body_start_index(doc: Document) -> int:
    """跳过封面：首个含分页符的段落之后为正文起点。"""
    body = doc.element.body
    kids = list(body)
    for idx, child in enumerate(kids):
        if child.tag != qn("w:p"):
            continue
        for br in child.iter(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return idx + 1
    return 0


def parse_docx(path: Path) -> dict:
    doc = Document(str(path))
    start = _body_start_index(doc)

    headings: list[tuple[int, str]] = []
    for p in doc.paragraphs:
        name = (p.style.name or "") if p.style else ""
        if name == "Heading 1":
            headings.append((1, _norm_heading(p.text)))
        elif name == "Heading 2":
            headings.append((2, _norm_heading(p.text)))
        elif name == "Heading 3":
            headings.append((3, _norm_heading(p.text)))
        elif name == "Heading 4":
            headings.append((4, _norm_heading(p.text)))

    inline_shapes = len(doc.inline_shapes)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    inline_shapes += len(p._element.findall(".//" + qn("a:blip")))

    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    swu_in_docx = sorted(set(_SWU.findall(full_text)))

    return {
        "headings": headings,
        "tables": len(doc.tables),
        "inline_images": inline_shapes,
        "paragraphs": len(doc.paragraphs),
        "swu_ids": swu_in_docx,
        "body_start": start,
        "full_text": full_text,
    }


def _check_function_tables(docx_path: Path) -> list[str]:
    """抽查 HTML 源表在 DOCX 中是否拆成独立行（避免合并到返回值格）。"""
    from docx import Document

    errors: list[str] = []
    d = Document(str(docx_path))
    for t in d.tables:
        blob = "\n".join(c.text for row in t.rows for c in row.cells)
        if "Drv_Ipcs_Core_Cmp" in blob and "IPCS_001" in blob:
            has_def = any("函数定义文件" in c.text for row in t.rows for c in row.cells)
            has_req = "满足需求" in blob
            merged_bad = "函数定义文件" in blob and not has_def
            if merged_bad or not has_req:
                errors.append(
                    "ipcsShmInit-style table: missing separate rows for 定义/声明/需求"
                )
            break
    return errors


def _docx_body_media_exts(docx_path: Path) -> dict[str, int]:
    """统计 word/media 中各扩展名数量（正文插图应为 SVG）。"""
    import zipfile

    exts: dict[str, int] = {}
    with zipfile.ZipFile(docx_path) as z:
        for name in z.namelist():
            if not name.startswith("word/media/"):
                continue
            ext = Path(name).suffix.lower() or "(none)"
            exts[ext] = exts.get(ext, 0) + 1
    return exts


def _aspect_ratio_mismatches(
    md_image_paths: list[str], docx_path: Path, tolerance: float = 0.02
) -> list[tuple[int, str, float, float]]:
    """返回 (序号, 文件名, viewBox_ratio, docx_ratio) 列表。"""
    doc = Document(str(docx_path))
    inlines = doc.element.body.findall(".//" + qn("wp:inline"))
    if len(inlines) != len(md_image_paths):
        return [(0, "count_mismatch", 0.0, 0.0)]
    bad: list[tuple[int, str, float, float]] = []
    for i, (rel, inline_el) in enumerate(zip(md_image_paths, inlines), start=1):
        svg_path = ROOT / rel
        if not svg_path.is_file():
            continue
        wh = parse_svg_viewbox_wh(svg_path.read_text(encoding="utf-8", errors="ignore"))
        if wh is None:
            continue
        exp = wh[1] / wh[0]
        ext = inline_el.find(qn("wp:extent"))
        cx = int(ext.get("cx") or 0) if ext is not None else 0
        cy = int(ext.get("cy") or 0) if ext is not None else 0
        act = cy / cx if cx else 0.0
        if exp > 0 and abs(act - exp) / exp > tolerance:
            bad.append((i, Path(rel).name, exp, act))
    return bad


def _non_svg_md_images(images: list[str]) -> list[str]:
    return [
        p
        for p in images
        if not p.startswith(("http://", "https://", "data:"))
        and not p.lower().endswith(".svg")
    ]


def _count_body_svg_blips(docx_path: Path) -> int:
    """正文内联插图通过 Word svgBlip 嵌入 SVG 的数量。"""
    from docx.oxml.ns import qn

    SVGNS = "http://schemas.microsoft.com/office/drawing/2016/SVG/main"
    RNS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    d = Document(str(docx_path))
    n = 0
    for blip in d.element.body.findall(".//" + qn("a:blip")):
        if blip.find(f".//{{{SVGNS}}}svgBlip") is not None:
            n += 1
    return n


def _count_html_style_tables(docx_path: Path) -> int:
    from docx import Document

    n = 0
    d = Document(str(docx_path))
    for t in d.tables:
        if len(t.columns) >= 5:
            blob = t.rows[0].cells[0].text if t.rows else ""
            if "对应软件架构" in blob or "对应软件架构ID" in blob:
                n += 1
    return n


def compare(md: dict, docx: dict, docx_path: Path) -> tuple[list[str], list[str]]:
    lines: list[str] = []
    errors: list[str] = []
    warns: list[str] = []

    lines.append("=== MD / DOCX consistency report ===\n")

    md_h = [h[1] for h in md["headings"] if h[0] <= 4]
    docx_h = [h[1] for h in docx["headings"]]
    lines.append(f"Markdown headings (h1-h4): {len(md_h)}")
    lines.append(f"DOCX Heading 1-4 paragraphs: {len(docx_h)}")

    md_set = set(md_h)
    docx_set = set(docx_h)
    only_md = sorted(md_set - docx_set)[:30]
    only_docx = sorted(docx_set - md_set)[:30]
    if only_md:
        warns.append(f"{len(md_set - docx_set)} heading(s) in MD not found in DOCX (sample): {only_md[:8]}")
    if len(docx_h) < len(md_h) * 0.85:
        errors.append(
            f"DOCX heading count ({len(docx_h)}) < 85% of MD ({len(md_h)})"
        )

    lines.append(f"\nMarkdown pipe tables (blocks): {md['tables']}")
    lines.append(f"Markdown HTML function tables: {md.get('html_tables', 0)}")
    html5 = _count_html_style_tables(docx_path)
    lines.append(f"DOCX HTML-style function tables (5+ cols): {html5}")
    lines.append(f"DOCX tables (total): {docx['tables']}")
    if md.get("html_tables", 0) and html5 < md["html_tables"] * 0.9:
        errors.append(
            f"HTML function tables in DOCX ({html5}) < 90% of MD HTML tables ({md['html_tables']})"
        )
    if docx["tables"] < (md["tables"] + md.get("html_tables", 0)) * 0.85:
        errors.append(
            f"DOCX tables ({docx['tables']}) < 85% of MD table blocks ({md['tables']})"
        )

    lines.append(f"\nMarkdown image refs: {len(md['images'])}")
    lines.append(f"DOCX inline images (blips): {docx['inline_images']}")
    non_svg = _non_svg_md_images(md["images"])
    if non_svg:
        errors.append(
            f"MD 插图必须为 .svg，发现 {len(non_svg)} 个非 SVG: {non_svg[:8]}"
        )
    media_exts = _docx_body_media_exts(docx_path)
    lines.append(f"DOCX word/media extensions: {media_exts}")
    svg_media = media_exts.get(".svg", 0)
    if docx["inline_images"] and svg_media < docx["inline_images"] * 0.5:
        warns.append(
            f"DOCX media 中 SVG 偏少 ({svg_media} svg vs {docx['inline_images']} blips)；"
            "若含封面 PNG 可忽略，正文插图应 SVG 直嵌"
        )
    if md["missing_image_files"]:
        errors.append(
            f"Missing image files on disk: {md['missing_image_files'][:10]}"
            + (f" ... (+{len(md['missing_image_files'])-10})" if len(md["missing_image_files"]) > 10 else "")
        )
    if docx["inline_images"] < len(md["images"]):
        errors.append(
            f"DOCX images ({docx['inline_images']}) < MD refs ({len(md['images'])}) — "
            "流程图/序列图可能未嵌入，检查 Pandoc resource-path 与 SVG 直嵌"
        )
    svg_blips = _count_body_svg_blips(docx_path)
    lines.append(f"DOCX body SVG (svgBlip): {svg_blips}")
    if svg_blips < len(md["images"]):
        errors.append(
            f"正文 SVG 嵌入不足: svgBlip={svg_blips}, MD refs={len(md['images'])}"
        )
    elif docx["inline_images"] < len(md["images"]) * 0.85:
        errors.append(
            f"DOCX images ({docx['inline_images']}) < 85% of MD refs ({len(md['images'])})"
        )

    aspect_bad = _aspect_ratio_mismatches(md["images"], docx_path)
    if aspect_bad:
        errors.append(
            f"SVG viewBox 宽高比与 DOCX extent 不一致: {len(aspect_bad)} 张，"
            f"例 #{aspect_bad[0][0]} {aspect_bad[0][1]} (viewBox={aspect_bad[0][2]:.3f} docx={aspect_bad[0][3]:.3f})"
        )
    else:
        lines.append(
            f"\nSVG aspect ratios (viewBox vs DOCX): {len(md['images'])}/{len(md['images'])} OK"
        )

    lines.append(f"\nSWU IDs in MD: {len(md['swu_ids'])}")
    lines.append(f"SWU IDs in DOCX: {len(docx['swu_ids'])}")
    missing_swu = sorted(set(md["swu_ids"]) - set(docx["swu_ids"]))
    if missing_swu:
        errors.append(f"SWU IDs in MD missing from DOCX: {missing_swu}")

    docx_text = docx["full_text"]
    key_phrases = [
        "5.7 RTOS",
        "6.7 Linux",
        "sequence diagram",
        "SWU_IPCS_LINUX_OS_UIO",
        "ipcsShmInit",
    ]
    lines.append("\nKey phrase spot-check:")
    for phrase in key_phrases:
        in_doc = phrase.lower() in docx_text.lower()
        status = "OK" if in_doc else "MISSING"
        if not in_doc and phrase not in ("rtos_seq_init", "linux_seq_uio_init"):
            warns.append(f"Phrase not in DOCX body: {phrase}")
        lines.append(f"  [{status}] {phrase}")

    lines.append(f"\nFunction subsections (### x.y.z name) in MD: {len(md['func_sections'])}")

    if warns:
        lines.append("\n--- WARNINGS ---")
        lines.extend(f"  [WARN] {w}" for w in warns)
    if errors:
        lines.append("\n--- ERRORS ---")
        lines.extend(f"  [ERROR] {e}" for e in errors)
    else:
        lines.append("\n--- RESULT: PASS (no ERROR) ---")

    lines.append(f"\nDOCX paragraphs total: {docx['paragraphs']}")
    return lines, errors


def main() -> None:
    docx = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else DOCX_PATH.resolve()
    if not MD_PATH.is_file():
        sys.exit(f"missing {MD_PATH}")
    if not docx.is_file():
        sys.exit(f"missing {docx}")

    md_data = parse_md(MD_PATH.read_text(encoding="utf-8"))
    docx_data = parse_docx(docx)
    report_lines, errors = compare(md_data, docx_data, docx)
    ft_errs = _check_function_tables(docx)
    if ft_errs:
        errors.extend(ft_errs)
        if not any("--- ERRORS ---" in ln for ln in report_lines):
            report_lines.append("\n--- ERRORS ---")
        report_lines.extend(f"  [ERROR] {e}" for e in ft_errs)
        if "--- RESULT: PASS (no ERROR) ---" in report_lines:
            report_lines.remove("--- RESULT: PASS (no ERROR) ---")
    report = "\n".join(report_lines)
    REPORT_PATH.write_text(report, encoding="utf-8")
    print(report)
    print(f"\n[written] {REPORT_PATH}")
    sys.exit(1 if errors else 0)


if __name__ == "__main__":
    main()
