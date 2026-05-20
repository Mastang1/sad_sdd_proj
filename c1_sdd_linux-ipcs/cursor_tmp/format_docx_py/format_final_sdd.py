# -*- coding: utf-8 -*-
r"""
final_sdd.docx（TF）批量版式整理脚本
====================================

依据工作区约定，TF 指：
``C:\tangyapeng\docs\StarGather\c1_sdd\final_sdd.docx``

本脚本仅执行 prompt-format.md 中列出的版式任务，不改动未要求的逻辑与内容语义。

**依赖**：``pip install python-docx``

**使用方式**（在项目根目录 `c1_sdd` 下）::

    python cursor_tmp/format_docx_py/format_final_sdd.py
    python cursor_tmp/format_docx_py/format_final_sdd.py "D:\\path\\final_sdd.docx"

仅执行任务 3、6（跳过其余任务）::

    python format_docx_py/format_tf_tasks_3_6.py
    python format_docx_py/format_tf_tasks_3_6.py "D:\\path\\final_sdd.docx"

仅执行任务 6（跳过其余任务）::

    python format_docx_py/format_tf_task_6.py
    python format_docx_py/format_tf_task_6.py "D:\\path\\final_sdd.docx"

仅执行任务 7（跳过其余任务）::

    python format_docx_py/format_tf_task_7.py
    python format_docx_py/format_tf_task_7.py "D:\\path\\final_sdd.docx"

直接修改 TF 原文件；若文件被 Word 占用导致无法保存，脚本将报错，请先关闭 Word 后重试。

**任务摘要**（与需求文档一一对应）：

1. （目录等）段落级 ``numPr`` 去手写号；（**Heading 1/2/3**）若样式在 ``styles.xml`` 中自带编号，
   则去掉正文里重复的章节号；再处理 ``Heading*``/``Title`` 中 **两段子写的连续编号**（粘贴残留）。
2. 为所有表格设置默认黑色实线边框（外框与内部网格）。
3. 将正文字段落格式统一为：微软雅黑、**11 号（11pt）**、不粗体；无首行缩进；单倍行距；段前段后 0 磅。
   不改标题、目录（Compact）、题注、图片说明等样式。结束后做标题完整性校验。
4. 所有表格单元格内文字：微软雅黑、10pt、不粗体。
5. 遍历内联插图：在 **不超过版心（四边页边距内可排版区域）** 的前提下，一般宽度设为 **14 cm**；
   **3.2 Files** 节内插图宽度为 **12 cm**；若目标宽度超出版心则先压至版心宽；若高度超出版心高度再同比缩小。
   全程 **保持原始宽高比**。
6. 主文档正文流中含插图/图表的段落（非表格单元格内）：**无缩进**、**居中对齐**（相对页面版心，而非单元格对齐）。
7. **3.2 Files**：删除 **3.2.1～3.2.18** 各插图**下一行**中形如 ``3.2.x …`` 的说明单行；**3.3 / 3.4**：删除插图下含 ``processing flow`` 的编号说明行，并在**每个插图前**插入正文格式的 ``processing flow`` 段。

版本：与 TF 配套维护。
"""

from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Emu, Pt
from docx.shape import InlineShape
from docx.table import Table

# ---------------------------------------------------------------------------
# 路径约定
# ---------------------------------------------------------------------------
_SCRIPT_DIR = Path(__file__).resolve().parent
_CURSOR_TMP = _SCRIPT_DIR.parent
if str(_CURSOR_TMP) not in sys.path:
    sys.path.insert(0, str(_CURSOR_TMP))
from workspace_paths import WORKSPACE_ROOT, FINAL_SDD_DOCX

DEFAULT_TF = FINAL_SDD_DOCX

# 视为「正文」的段落样式（仅这些会应用任务 3）
BODY_STYLE_NAMES = frozenset(
    {
        "Normal",
        "First Paragraph",
        "Body Text",
        "Block Text",
    }
)

# 题注、标题等不参与任务 3
SKIP_BODY_STYLES = frozenset(
    {
        "Compact",
        "Image Caption",
        "Captioned Figure",
    }
)

_WESTERN_FONT = "Microsoft YaHei"
_EAST_ASIA_FONT = "微软雅黑"
_BODY_FONT_PT = 11  # 正文 11 号（与 prompt-format 任务 3 一致）
_TABLE_FONT_PT = 10
# 任务 5：目标宽度（厘米）；3.2 Files 节单独为 12 cm
_DEFAULT_FIGURE_WIDTH_CM = 14.0
_FILES_SECTION32_FIGURE_WIDTH_CM = 12.0
_EMU_PER_CM = 360000

# ---------------------------------------------------------------------------
# 标题样式判断
# ---------------------------------------------------------------------------
def _is_outline_heading_paragraph(paragraph) -> bool:
    """
    是否属于「文档大纲标题」段落（任务 1 中去双重章节号作用域）。

    :param paragraph: ``docx.text.paragraph.Paragraph``
    :return: 为 ``Heading 1``～``Heading 9`` 或 ``Title`` 时返回 True
    """
    name = paragraph.style.name if paragraph.style else ""
    if name == "Title":
        return True
    if name.startswith("Heading"):
        return True
    return False


def _tf_path() -> Path:
    if len(sys.argv) > 1:
        return Path(sys.argv[1]).resolve()
    return DEFAULT_TF.resolve()


def _heading_style_has_numbering(document: Document, style_name: str) -> bool:
    """
    判断段落样式是否在 ``styles.xml`` 中定义了 ``w:numPr``（大纲编号由样式呈现）。

    :param document: 当前文档
    :param style_name: 如 ``Heading 2``
    :return: 样式含 ``numPr`` 时为 True
    """
    try:
        st = document.styles[style_name]
    except KeyError:
        return False
    ppr = st.element.pPr
    return ppr is not None and ppr.numPr is not None


def _safe_strip_outline_prefix_once(text: str) -> tuple[str, bool]:
    """
    去掉字符串左端一层章节号前缀（``1 `` / ``3.2 `` / ``3.3.1 ``）。

    若前缀为 **4 位及以上且全数字无点**（如年份 ``2024 ``）则不删除，避免误伤正文。

    :param text: 待处理文本
    :return: (新文本, 是否删除了前缀)
    """
    s = text.lstrip()
    if not s:
        return text, False
    m = re.match(r"^(\d+(?:\.\d+)*)\s+", s)
    if not m:
        return text, False
    tok = m.group(1)
    if "." not in tok and len(tok) >= 4 and tok.isdigit():
        return text, False
    rest = s[m.end() :]
    return rest, True


def strip_heading_manual_numbers_for_style_numbering(document: Document, paragraphs: Iterable) -> int:
    """
    删除 **Heading 1/2/3** 中与「样式自带编号」重复的手写章节号。

    TF 中 ``heading 1``～``heading 3`` 均在 ``styles.xml`` 内嵌 ``w:numPr``，正文 run 若再含 ``1 ``、
    ``3.2 `` 等，会出现两套编号。本函数对同一段连续剥离最多 **3** 层合法前缀。

    :param document: Document
    :param paragraphs: ``document.paragraphs``
    :return: 修改的标题段落数量
    """
    changed = 0
    for lvl in (1, 2, 3):
        name = f"Heading {lvl}"
        for para in paragraphs:
            if (para.style.name if para.style else "") != name:
                continue
            if not _heading_style_has_numbering(document, name):
                continue
            if para._p.xpath(".//w:drawing"):
                continue
            t0 = para.text or ""
            if not t0.strip():
                continue
            t = t0
            any_strip = False
            for _ in range(3):
                t2, did = _safe_strip_outline_prefix_once(t)
                if not did:
                    break
                t = t2
                any_strip = True
            if any_strip and t != t0:
                para.text = t
                changed += 1
    return changed


# ---------------------------------------------------------------------------
# 任务 1：列表编号与手写章节号重复
# ---------------------------------------------------------------------------
def remove_duplicate_numbering_after_list(paragraphs: Iterable) -> int:
    """
    对带 ``w:numPr`` 的段落：若段内文本以「1 / 1.2 / 1.2.3」形式章节号开头，则删除该手写前缀，
    避免与 Word 自动编号叠加成双编号。

    跳过含 ``w:drawing`` 的段落，避免破坏图文混排。

    :param paragraphs: 一般为 ``document.paragraphs``
    :return: 实际修改的段落数量
    """
    changed = 0
    pat = re.compile(r"^\d+(?:\.\d+)*\s+")
    for para in paragraphs:
        p = para._p
        ppr = p.pPr
        if ppr is None or ppr.numPr is None:
            continue
        if p.xpath(".//w:drawing"):
            continue
        full = para.text or ""
        if not full.strip():
            continue
        new_t = pat.sub("", full, count=1)
        if new_t == full:
            continue
        para.text = new_t
        changed += 1
    return changed


def remove_duplicate_heading_section_numbers(paragraphs: Iterable) -> int:
    """
    任务 1（补充）：处理 **各级标题**（``Heading*`` / ``Title``）中因拷贝遗留的重复章节号。

    仅改段落**纯文本**拼接结果（``paragraph.text``），不处理域代码；跳过含 ``w:drawing`` 的段落。

    对 **两个连续「纯章节号」词元**（如 ``4.3 3.3 外部接口``）：删第二段，保留首段与余下正文
    （比单纯正则更稳，不依赖 Tab）。

    仍保留上述正则，用于 **Tab 分隔**、**同号重复** 等边角情况。

    :param paragraphs: ``document.paragraphs``
    :return: 修改的标题段落数
    """
    changed = 0
    num_token = re.compile(r"^\d+(?:\.\d+)*$")
    two_prefix = re.compile(
        r"^(\d+(?:\.\d+)*)([\s\u00a0\u3000]+)(\d+(?:\.\d+)*)([\s\u00a0\u3000]+)(.+)$"
    )
    tab_between = re.compile(r"^(\d+(?:\.\d+)*)\t+(\d+(?:\.\d+)*)\s*(.+)$")
    repeat_same = re.compile(r"^(\d+(?:\.\d+)*\s+)\1+(.+)$")

    for para in paragraphs:
        if not _is_outline_heading_paragraph(para):
            continue
        if para._p.xpath(".//w:drawing"):
            continue
        t = para.text or ""
        if not t.strip():
            continue
        new_t = t
        # 前两个词均为 1 / 1.2 / 1.2.3 形式 → 去掉第二段
        parts = t.strip().split()
        if len(parts) >= 3 and num_token.match(parts[0]) and num_token.match(parts[1]):
            new_t = f"{parts[0]} {' '.join(parts[2:])}".strip()
        else:
            m = two_prefix.match(t.strip())
            if m:
                new_t = f"{m.group(1)} {m.group(5).lstrip()}"
            else:
                m = tab_between.match(t.strip())
                if m:
                    new_t = f"{m.group(1)} {m.group(3)}"
                else:
                    m = repeat_same.match(t.strip())
                    if m:
                        new_t = m.group(1) + m.group(2)
        if new_t != t:
            para.text = new_t
            changed += 1
    return changed


# ---------------------------------------------------------------------------
# 任务 2：表格实线边框
# ---------------------------------------------------------------------------
def apply_solid_table_borders(table: Table) -> None:
    """
    为单个表格设置 ``tblBorders``：top/left/bottom/right/insideH/insideV 均为单线（single）、默认线宽与颜色。

    若已存在 ``w:tblBorders`` 则先移除再写入，保证整表一致。

    :param table: ``python-docx`` 的 Table 对象
    """
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def apply_all_table_borders(document: Document) -> int:
    """
    遍历文档主故事中全部表格并设置实线边框。

    :param document: 已打开的 Document
    :return: 处理的表格数量
    """
    n = 0
    for tbl in document.tables:
        apply_solid_table_borders(tbl)
        n += 1
    return n


# ---------------------------------------------------------------------------
# 字体辅助（东亚字体在 OOXML 中需 rFonts/@eastAsia）
# ---------------------------------------------------------------------------
def _set_run_font(
    run,
    *,
    western: str,
    east_asia: str,
    size_pt: int,
    bold: bool,
) -> None:
    """
    设置单个 run 的西文/东亚字体、字号、粗体；写 ``w:rFonts`` 以兼容 Word 中文显示。

    :param run: ``docx.text.run.Run``
    :param western: 西文字体名（如 Microsoft YaHei）
    :param east_asia: 东亚字体名（如 微软雅黑）
    :param size_pt: 磅值
    :param bold: 是否粗体
    """
    run.font.name = western
    run.font.size = Pt(size_pt)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), western)
    r_fonts.set(qn("w:hAnsi"), western)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def format_body_paragraphs(document: Document) -> int:
    """
    任务 3：仅对正文样式段落设置段格式与字体（微软雅黑、11 pt、不粗体、无首行缩进、单倍行距、段前后 0）。

    标题（Heading* / Title）、目录（Compact）、题注等一律跳过。

    :param document: Document
    :return: 修改过的正文段落数
    """
    n = 0
    for para in document.paragraphs:
        name = para.style.name if para.style else ""
        if name.startswith("Heading") or name in ("Title", "Subtitle"):
            continue
        if name in SKIP_BODY_STYLES:
            continue
        if name not in BODY_STYLE_NAMES:
            continue
        pf = para.paragraph_format
        pf.first_line_indent = Pt(0)
        pf.left_indent = Pt(0)
        pf.right_indent = Pt(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.line_spacing = None
        for run in para.runs:
            if not run._element.xpath(".//w:drawing"):
                _set_run_font(
                    run,
                    western=_WESTERN_FONT,
                    east_asia=_EAST_ASIA_FONT,
                    size_pt=_BODY_FONT_PT,
                    bold=False,
                )
        n += 1
    return n


def format_figure_chart_paragraphs_layout(document: Document) -> int:
    """
    任务 6：将**主文档正文流**中含插图或图表的段落设置为**行首无缩进**（左右缩进清零）、**居中对齐**。

    仅处理 ``document.paragraphs`` 所含段落（与 python-docx 一致：不含表格单元格内段落，
    避免改变「单元格内对齐」）。现代 Word 中图表、图片等多在 ``w:drawing`` 下；
    旧式 VML 图像可能在 ``w:pict`` 中，一并识别。

    不修改标题类段落（Heading* / Title / Subtitle），以免误动大纲标题行。

    :param document: 已打开的 Document（主文档部件）
    :return: 被设置段落格式的段落数量
    """
    n = 0
    for para in document.paragraphs:
        name = para.style.name if para.style else ""
        if name.startswith("Heading") or name in ("Title", "Subtitle"):
            continue
        el = para._element
        has_drawing = bool(el.findall(".//" + qn("w:drawing")))
        has_pict = bool(el.findall(".//" + qn("w:pict")))
        if not (has_drawing or has_pict):
            continue
        pf = para.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Pt(0)
        pf.left_indent = Pt(0)
        pf.right_indent = Pt(0)
        n += 1
    return n


def verify_task6_figure_paragraph_layout(document: Document) -> tuple[bool, list[str], dict]:
    """
    核对任务 6：主文档流（``document.paragraphs``）中含图/表的段落须**居中**且**首行/左/右缩进为 0**（磅容差 0.05）。

    统计表格单元格内含图段落数（任务 6 按需求不调整单元格内对齐，仅供报告）。

    :param document: Document（建议为保存后重新加载，以排除未落盘差异）
    :return: ``(全部通过, 问题列表, 统计字典)``
    """
    issues: list[str] = []
    stats: dict = {
        "main_story_figure_paragraphs": 0,
        "main_story_passed": 0,
        "table_cell_figure_paragraphs": 0,
    }
    for para in document.paragraphs:
        name = para.style.name if para.style else ""
        if name.startswith("Heading") or name in ("Title", "Subtitle"):
            continue
        el = para._element
        has_fig = bool(el.findall(".//" + qn("w:drawing"))) or bool(
            el.findall(".//" + qn("w:pict"))
        )
        if not has_fig:
            continue
        stats["main_story_figure_paragraphs"] += 1
        seq = stats["main_story_figure_paragraphs"]
        pf = para.paragraph_format
        ok = True
        if pf.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            issues.append(
                f"主文档插图段 #{seq}：段落对齐为 {pf.alignment!r}，要求 CENTER。"
            )
            ok = False
        for label, attr in (
            ("first_line_indent", pf.first_line_indent),
            ("left_indent", pf.left_indent),
            ("right_indent", pf.right_indent),
        ):
            if attr is None:
                continue
            try:
                pt = attr.pt
            except AttributeError:
                continue
            if abs(pt) > 0.05:
                issues.append(
                    f"主文档插图段 #{seq}：{label}={pt:g} pt，要求 0。"
                )
                ok = False
        if ok:
            stats["main_story_passed"] += 1

    for p_el in _iter_body_paragraphs_in_order(document.element.body):
        parent = p_el.getparent()
        if parent is None or parent.tag != qn("w:tc"):
            continue
        if p_el.findall(".//" + qn("w:drawing")) or p_el.findall(
            ".//" + qn("w:pict")
        ):
            stats["table_cell_figure_paragraphs"] += 1

    return (len(issues) == 0, issues, stats)


def format_all_table_cells(document: Document) -> int:
    """
    任务 4：遍历所有表格单元格内全部 run（不含直接嵌入的形状文本），字体微软雅黑 10pt、不粗体。

    :param document: Document
    :return: 有 run 被格式化的单元格数量（每个单元格最多计 1 次）
    """
    cells_touched = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                hit = False
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run._element.xpath(".//w:drawing"):
                            continue
                        _set_run_font(
                            run,
                            western=_WESTERN_FONT,
                            east_asia=_EAST_ASIA_FONT,
                            size_pt=_TABLE_FONT_PT,
                            bold=False,
                        )
                        hit = True
                if hit:
                    cells_touched += 1
    return cells_touched


# ---------------------------------------------------------------------------
# 任务 5：内联插图 — 版心（页边距内）约束 + 目标宽度 + 等比例
# ---------------------------------------------------------------------------
def _max_printable_box_emu(document: Document) -> tuple[int, int]:
    """
    各节版心宽、高分别取全文档最小值（EMU），保证任一节的页边距下插图都不越界。
    docxcompose 合并后 ``document.sections`` 可能为空，此时回退 A4 + 2.54 cm 页边距。
    """
    min_cx: int | None = None
    min_cy: int | None = None
    for s in document.sections:
        if s.page_width is None or s.left_margin is None or s.right_margin is None:
            continue
        if s.page_height is None or s.top_margin is None or s.bottom_margin is None:
            continue
        cw = int(s.page_width - s.left_margin - s.right_margin)
        ch = int(s.page_height - s.top_margin - s.bottom_margin)
        if cw > 0:
            min_cx = cw if min_cx is None else min(min_cx, cw)
        if ch > 0:
            min_cy = ch if min_cy is None else min(min_cy, ch)
    if min_cx is None or min_cx <= 0:
        min_cx = int(Cm(15.92))
    if min_cy is None or min_cy <= 0:
        min_cy = int(Cm(22.0))
    return min_cx, min_cy


def _iter_p_in_table(tbl_el) -> Iterable:
    """深度优先遍历表格内所有 ``w:p``（含嵌套表）。"""
    for tr in tbl_el.findall(qn("w:tr")):
        for tc in tr.findall(qn("w:tc")):
            for elem in tc:
                if elem.tag == qn("w:p"):
                    yield elem
                elif elem.tag == qn("w:tbl"):
                    yield from _iter_p_in_table(elem)


def _iter_body_paragraphs_in_order(body_el) -> Iterable:
    """body 子级顺序遍历 ``w:p`` 与表内段落（与 Word 阅读顺序一致）。"""
    for child in body_el:
        if child.tag == qn("w:p"):
            yield child
        elif child.tag == qn("w:tbl"):
            yield from _iter_p_in_table(child)


def _paragraph_style_display_name(document: Document, p_el) -> str | None:
    """由段落 XML 的 ``pStyle`` 解析样式显示名（如 ``Heading 2``）。"""
    p_pr = p_el.pPr
    if p_pr is None or p_pr.pStyle is None:
        return None
    sid = p_pr.pStyle.val
    try:
        return document.styles.get_by_id(sid, WD_STYLE_TYPE.PARAGRAPH).name
    except (KeyError, ValueError):
        return None


def _plain_text_from_p_el(p_el) -> str:
    """段落内纯文本（``w:t`` + 制表符），用于识别标题文案。"""
    parts: list[str] = []
    for node in p_el.iter():
        if node.tag == qn("w:t"):
            parts.append(node.text or "")
        elif node.tag == qn("w:tab"):
            parts.append("\t")
    return "".join(parts)


# ---------------------------------------------------------------------------
# 任务 7：3.2 图下编号说明行删除；3.3/3.4 processing flow 行前移为图前正文段
# ---------------------------------------------------------------------------
_RE_32_FIGURE_NEXT_CAPTION = re.compile(r"^\s*3\.2\.(?:1[0-8]|[1-9])\b")


def _h2_ends_section_33_34(text: str) -> bool:
    """
    判断 ``Heading 2`` 是否表示 3.3/3.4 叙述结束（进入后续章节，如全局变量）。

    TF 中存在拼写 ``Gobal variants``，按子串兼容匹配。

    :param text: 标题段落全文
    :return: 若为结束边界标题返回 True
    """
    s = text.strip().casefold()
    if "gobal variant" in s:
        return True
    if "global variant" in s:
        return True
    return False


def _task7_regions_at_index(document: Document, all_p_els: list, idx: int) -> tuple[bool, bool]:
    """
    根据 ``all_p_els[0:idx]`` 中已出现的 ``Heading 2``，判定当前索引是否在 3.2 Files 或 3.3/3.4 范围。

    :param document: Document
    :param all_p_els: ``_iter_body_paragraphs_in_order`` 快照
    :param idx: 当前段落索引
    :return: ``(in_32_files, in_33_34)``
    """
    in_32 = False
    in_33_34 = False
    for j in range(idx):
        p_el = all_p_els[j]
        st = _paragraph_style_display_name(document, p_el)
        if st != "Heading 2":
            continue
        tx = _plain_text_from_p_el(p_el)
        s = tx.strip().casefold()
        if s == "files":
            in_32, in_33_34 = True, False
        elif "external interfaces" in s:
            in_32, in_33_34 = False, True
        elif _h2_ends_section_33_34(tx):
            in_33_34 = False
    return in_32, in_33_34


def _p_el_has_figure_or_chart(p_el) -> bool:
    """段落是否含插图/图表（``w:drawing`` 或旧式 ``w:pict``）。"""
    return bool(
        p_el.findall(".//" + qn("w:drawing"))
        or p_el.findall(".//" + qn("w:pict"))
    )


def _is_33_34_processing_flow_caption(text: str) -> bool:
    """
    是否为目标说明行：行内含 ``processing flow``，且行首为 ``3.3.x`` / ``3.4.x`` 编号。

    :param text: 插图**后**紧邻段落全文
    :return: 是否匹配
    """
    t = (text or "").strip()
    if "processing flow" not in t.casefold():
        return False
    return bool(re.match(r"^\s*3\.(?:3|4)\.\d+", t))


def _cell_from_tc(document: Document, tc_el):
    """
    由 ``w:tc`` 元素定位 ``docx.table.Cell``，用于在单元格内插入段时构造 ``Paragraph`` 容器。

    :param document: Document
    :param tc_el: ``w:tc`` 元素
    :return: 单元格或 None
    """
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell._tc is tc_el:
                    return cell
    return None


def _paragraph_block_container(document: Document, p_el):
    """
    解析 ``w:p`` 所在块容器（``body`` 或表格单元格），供 ``Paragraph(xml, container)`` 使用。

    :param document: Document
    :param p_el: ``w:p`` 元素
    :return: ``document._body`` 或 ``Cell``
    :raises ValueError: 父级既不是 body 也不是 tc
    """
    parent = p_el.getparent()
    if parent is None:
        raise ValueError("段落已从树中移除")
    if parent.tag == qn("w:body"):
        return document._body
    if parent.tag == qn("w:tc"):
        cell = _cell_from_tc(document, parent)
        if cell is None:
            raise ValueError("无法将 w:tc 映射到 python-docx Cell")
        return cell
    raise ValueError(f"不支持的段落父级标签: {parent.tag}")


def _apply_task7_body11_paragraph_format(paragraph) -> None:
    """
    任务 7 插入段落的段格式：与任务 3 正文一致（无缩进、单倍行距、段前后 0）。

    :param paragraph: ``Paragraph``
    """
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = None


def _remove_p_element(p_el) -> None:
    """从 OOXML 树中删除 ``w:p``。"""
    parent = p_el.getparent()
    if parent is not None:
        parent.remove(p_el)


def _insert_processing_flow_paragraph_before(document: Document, p_fig_el) -> None:
    """
    在插图段落 ``p_fig_el`` **之前**插入单独一段，文本为 ``processing flow``，
    字体与段落格式同任务 3 正文（微软雅黑 11 pt、不粗体）。

    :param document: Document
    :param p_fig_el: 含 ``w:drawing`` 的 ``w:p`` 元素
    """
    container = _paragraph_block_container(document, p_fig_el)
    new_el = OxmlElement("w:p")
    p_fig_el.addprevious(new_el)
    para = Paragraph(new_el, container)
    try:
        para.style = document.styles["Normal"]
    except KeyError:
        pass
    _apply_task7_body11_paragraph_format(para)
    run = para.add_run("processing flow")
    _set_run_font(
        run,
        western=_WESTERN_FONT,
        east_asia=_EAST_ASIA_FONT,
        size_pt=int(_BODY_FONT_PT),
        bold=False,
    )


def apply_task7_caption_and_processing_flow(document: Document) -> dict[str, int]:
    """
    任务 7（``prompt-format.md``）：

    #. **3.2**（``Heading 2`` Files 起至 External Interfaces 止）：对每个含图段落，若**下一段落**
       全文匹配 ``3.2.1``～``3.2.18`` 开头的单行说明，则**删除该说明段**（不删图、不改其他字）。
    #. **3.3 / 3.4**（External Interfaces 之后至 ``Gobal variants`` / ``Global variants`` 类 ``Heading 2``
       之前）：对每个含图段落，若**下一段**为含 ``processing flow`` 的 ``3.3.*`` / ``3.4.*`` 说明行，
       则**先**在图段前插入正文样式 ``processing flow``，**再删除**原说明行。

    遍历顺序与 ``_iter_body_paragraphs_in_order`` 一致（含表内段）。仅修改上述段落，不触及其他内容。

    :param document: TF ``Document``
    :return: ``{\"n_delete_32\": int, \"n_move_processing_flow\": int}``
    """
    all_p = list(_iter_body_paragraphs_in_order(document.element.body))
    del_32_targets: list = []
    move_pf_jobs: list[tuple] = []
    for i, p_el in enumerate(all_p):
        if not _p_el_has_figure_or_chart(p_el):
            continue
        in_32, in_33_34 = _task7_regions_at_index(document, all_p, i)
        if i + 1 >= len(all_p):
            continue
        p_next = all_p[i + 1]
        tx_next = _plain_text_from_p_el(p_next).strip()
        if in_32 and _RE_32_FIGURE_NEXT_CAPTION.match(tx_next):
            del_32_targets.append(p_next)
        elif in_33_34 and _is_33_34_processing_flow_caption(tx_next):
            move_pf_jobs.append((p_el, p_next))
    for cap_el in del_32_targets:
        if cap_el.getparent() is not None:
            _remove_p_element(cap_el)
    for p_fig, p_cap in move_pf_jobs:
        if p_fig.getparent() is None or p_cap.getparent() is None:
            continue
        _insert_processing_flow_paragraph_before(document, p_fig)
        _remove_p_element(p_cap)
    return {
        "n_delete_32": len(del_32_targets),
        "n_move_processing_flow": len(move_pf_jobs),
    }


def _update_files_section32_region(style_name: str | None, text: str, in_files32: bool) -> bool:
    """
    根据 ``Heading 2`` 文本切换「3.2 Files」区间：从 ``Files`` 起至 ``External Interfaces`` 之前。

    :param style_name: 段落样式名
    :param text: 段落全文
    :param in_files32: 当前是否在 Files 小节内
    :return: 更新后的区间标志
    """
    if style_name != "Heading 2":
        return in_files32
    stripped = (text or "").strip()
    if stripped.casefold() == "files":
        return True
    if "external interfaces" in stripped.casefold():
        return False
    return in_files32


def _read_inline_natural_extent(inline_el, document: Document | None = None) -> tuple[int, int]:
    """委托 ``svg_extent_utils``：用内嵌 SVG viewBox 修正 Pandoc 错误 extent。"""
    from svg_extent_utils import read_inline_natural_extent

    return read_inline_natural_extent(inline_el, document)


def _read_inline_display_extent(inline_el, document: Document | None = None) -> tuple[int, int]:
    """当前显示 extent（与 natural 相同接口，用于校验）。"""
    return _read_inline_natural_extent(inline_el, document)


def _compute_target_extents(
    orig_cx: int,
    orig_cy: int,
    wanted_cx_emu: int,
    max_cx: int,
    max_cy: int,
) -> tuple[int, int]:
    """
    在满足版心宽高前提下，将宽度尽量设为 ``wanted_cx_emu``，并按原始宽高同比缩放高度；
    若高度仍超 ``max_cy``，再按高度上限回缩宽度。

    :param orig_cx: 原始宽度 EMU
    :param orig_cy: 原始高度 EMU
    :param wanted_cx_emu: 目标宽度（14 cm 或 12 cm）
    :param max_cx: 版心最大宽 EMU
    :param max_cy: 版心最大高 EMU
    """
    if orig_cx <= 0:
        return 1, 1
    aspect = orig_cy / float(orig_cx)
    cx = min(int(wanted_cx_emu), max_cx)
    cy = int(round(cx * aspect))
    if cy > max_cy:
        cy = max_cy
        cx = max(1, int(round(cy / aspect)))
    return cx, cy


def _apply_shape_extents(shape: InlineShape, cx_emu: int, cy_emu: int) -> None:
    """
    写回内联图的显示尺寸：同步 ``wp:extent`` 与 ``pic:spPr`` 变换区，并用 ``Emu`` 走属性适配接口。

    :param shape: ``InlineShape``
    :param cx_emu: 目标宽度（EMU，整型）
    :param cy_emu: 目标高度（EMU）
    """
    cx_i = max(1, int(cx_emu))
    cy_i = max(1, int(cy_emu))
    inl = shape._inline
    inl.extent.cx = cx_i
    inl.extent.cy = cy_i
    try:
        pic = inl.graphic.graphicData.pic
        pic.spPr.cx = cx_i
        pic.spPr.cy = cy_i
    except AttributeError:
        pass
    shape.width = Emu(cx_i)
    shape.height = Emu(cy_i)


def scale_inline_shapes_margin_and_width(document: Document) -> tuple[int, list[dict]]:
    """
    任务 5：按版心缩放所有 ``wp:inline``：一般宽 **14 cm**，**3.2 Files** 段内 **12 cm**；
    超出版心宽或高时先满足版心；全程保持原始宽高比。

    :param document: Document
    :return: ``(处理张数, 每张明细 dict 列表)``
    """
    max_cx, max_cy = _max_printable_box_emu(document)
    in_files32 = False
    reports: list[dict] = []
    n = 0
    wp_inline = qn("wp:inline")

    for p_el in _iter_body_paragraphs_in_order(document.element.body):
        st = _paragraph_style_display_name(document, p_el)
        tx = _plain_text_from_p_el(p_el)
        in_files32 = _update_files_section32_region(st, tx, in_files32)
        target_cm = (
            _FILES_SECTION32_FIGURE_WIDTH_CM if in_files32 else _DEFAULT_FIGURE_WIDTH_CM
        )
        wanted_emu = int(Cm(target_cm))

        for inline_el in p_el.findall(".//" + wp_inline):
            shape = InlineShape(inline_el)
            oc, oh = _read_inline_natural_extent(inline_el, document)
            if oc <= 0 or oh <= 0:
                continue
            exp_cx, exp_cy = _compute_target_extents(oc, oh, wanted_emu, max_cx, max_cy)
            _apply_shape_extents(shape, exp_cx, exp_cy)
            n += 1
            reports.append(
                {
                    "index": n,
                    "section_3_2_files": in_files32,
                    "target_width_cm": target_cm,
                    "printable_max_cm_w": max_cx / _EMU_PER_CM,
                    "printable_max_cm_h": max_cy / _EMU_PER_CM,
                    "before_cx_emu": oc,
                    "before_cy_emu": oh,
                    "after_cx_emu": exp_cx,
                    "after_cy_emu": exp_cy,
                    "after_width_cm": exp_cx / _EMU_PER_CM,
                    "after_height_cm": exp_cy / _EMU_PER_CM,
                }
            )

    return n, reports


def verify_inline_shapes_task5(
    document: Document, reports: list[dict]
) -> tuple[bool, list[str]]:
    """
    保存后重新打开文档，按与任务 5 相同遍历顺序核对每张图：版心、宽高比、与算法预期 extent 一致。

    :param document: 已保存的 TF 再加载
    :param reports: ``scale_inline_shapes_margin_and_width`` 返回的明细（含改前尺寸）
    :return: ``(是否全部通过, 问题描述列表)``
    """
    max_cx, max_cy = _max_printable_box_emu(document)
    issues: list[str] = []
    rep_iter = iter(reports)
    wp_inline = qn("wp:inline")

    in_files32 = False
    for p_el in _iter_body_paragraphs_in_order(document.element.body):
        st = _paragraph_style_display_name(document, p_el)
        tx = _plain_text_from_p_el(p_el)
        in_files32 = _update_files_section32_region(st, tx, in_files32)
        target_cm = (
            _FILES_SECTION32_FIGURE_WIDTH_CM if in_files32 else _DEFAULT_FIGURE_WIDTH_CM
        )
        wanted_emu = int(Cm(target_cm))

        for inline_el in p_el.findall(".//" + wp_inline):
            ac, ah = _read_inline_display_extent(inline_el, document)
            if ac <= 0 or ah <= 0:
                continue
            try:
                r = next(rep_iter)
            except StopIteration:
                issues.append("验证时插图数量多于处理报告条目。")
                break
            oc = r["before_cx_emu"]
            oh = r["before_cy_emu"]
            exp_cx, exp_cy = _compute_target_extents(oc, oh, wanted_emu, max_cx, max_cy)
            if abs(ac - exp_cx) > 2 or abs(ah - exp_cy) > 2:
                issues.append(
                    f"图 #{r['index']} extent 与算法预期不一致："
                    f"实际({ac},{ah}) 预期({exp_cx},{exp_cy})。"
                )
            if ac > max_cx + 2:
                issues.append(
                    f"图 #{r['index']} 宽度 {ac} EMU 超出版心最大宽 {max_cx} EMU。"
                )
            if ah > max_cy + 2:
                issues.append(
                    f"图 #{r['index']} 高度 {ah} EMU 超出版心最大高 {max_cy} EMU。"
                )
            tol = max(oc, oh, ac, ah, 1)
            if abs(ac * oh - ah * oc) > tol * 2:
                issues.append(
                    f"图 #{r['index']} 缩放后宽高比与原始不一致（"
                    f"原始 {oc}x{oh}，当前 {ac}x{ah}）。"
                )
    try:
        next(rep_iter)
        issues.append("验证结束仍有多余的报告条目（插图数量少于报告）。")
    except StopIteration:
        pass

    return (len(issues) == 0, issues)


# ---------------------------------------------------------------------------
# 标题完整性检测
# ---------------------------------------------------------------------------
def snapshot_heading_state(document: Document) -> list[tuple[str, str]]:
    """
    记录所有标题类段落（Heading 1/2/3 与 Title）的 (样式名, 全文)，用于前后对比。

    :param document: Document
    :return: 列表副本
    """
    out: list[tuple[str, str]] = []
    for para in document.paragraphs:
        name = para.style.name if para.style else ""
        if name.startswith("Heading") or name == "Title":
            out.append((name, para.text or ""))
    return out


def assert_heading_integrity(before: list[tuple[str, str]], after: list[tuple[str, str]]) -> None:
    """
    校验标题条数与文本是否与快照一致；不一致则抛出 AssertionError。

    :param before: ``snapshot_heading_state`` 结果
    :param after: 处理后的快照
    :raises AssertionError: 标题被意外改动
    """
    if before != after:
        for i, (a, b) in enumerate(zip(before, after)):
            if a != b:
                raise AssertionError(f"标题完整性失败：索引 {i}，修改前={a!r} 修改后={b!r}")
        if len(before) != len(after):
            raise AssertionError(
                f"标题条数不一致：before={len(before)} after={len(after)}"
            )


def remove_pandoc_placeholder_image_captions(document: Document) -> int:
    """删除 Pandoc 遗留的 ``imageN.png`` 占位题注（真图已在紧邻上一段）。"""
    pat = re.compile(r"^image\d+\.png$", re.I)
    body = document.element.body
    removed = 0
    for p_el in list(body.findall(qn("w:p"))):
        if _paragraph_style_display_name(document, p_el) != "Image Caption":
            continue
        if pat.match(_plain_text_from_p_el(p_el).strip()):
            body.remove(p_el)
            removed += 1
    return removed


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
def main() -> None:
    """打开 TF，依次执行任务 1～7 与标题校验，并保存。"""
    tf = _tf_path()
    if not tf.is_file():
        print(f"未找到 TF：{tf}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(tf))

    n_dup_list = remove_duplicate_numbering_after_list(doc.paragraphs)
    n_strip_heading = strip_heading_manual_numbers_for_style_numbering(doc, doc.paragraphs)
    n_dup_heading = remove_duplicate_heading_section_numbers(doc.paragraphs)
    # 标题在经过「目录手写号 + 标题双重号」规整后建立基准，其后任务不得再改标题文本
    headings_ref = snapshot_heading_state(doc)

    n_tbl_border = apply_all_table_borders(doc)
    n_tbl_cells = format_all_table_cells(doc)
    n_body = format_body_paragraphs(doc)
    n_shape, shape_reports = scale_inline_shapes_margin_and_width(doc)
    n_fig_layout = format_figure_chart_paragraphs_layout(doc)
    stats7 = apply_task7_caption_and_processing_flow(doc)
    n_img_cap = remove_pandoc_placeholder_image_captions(doc)

    headings_after = snapshot_heading_state(doc)
    assert_heading_integrity(headings_ref, headings_after)

    doc.save(str(tf))
    doc_check = Document(str(tf))
    task5_ok, task5_issues = verify_inline_shapes_task5(doc_check, shape_reports)

    pcw, pch = _max_printable_box_emu(doc_check)
    margin_line = (
        f"\n  版心（最紧节）最大宽高(cm): {pcw / _EMU_PER_CM:.4f} x "
        f"{pch / _EMU_PER_CM:.4f}"
    )
    print(
        "处理完成并已保存：",
        tf,
        f"\n  任务1a 列表段手写编号去除: {n_dup_list}",
        f"\n  任务1b Heading1-3 样式编号重复手写去除: {n_strip_heading}",
        f"\n  任务1c 标题双重章节号(Tab/两段子写): {n_dup_heading}",
        f"\n  任务2 表格边框: {n_tbl_border}",
        f"\n  任务3 正文段落: {n_body}",
        f"\n  任务4 表格单元格（计有文本的单元格）: {n_tbl_cells}",
        f"\n  任务5 缩放插图: {n_shape} 张",
        f"\n  任务6 插图/图表段落居中无缩进: {n_fig_layout}",
        f"\n  任务7 删3.2图下说明行: {stats7['n_delete_32']}；"
        f"3.3/3.4 processing flow 前移: {stats7['n_move_processing_flow']}",
        f"删除 imageN.png 占位题注: {n_img_cap} 行",
        "\n  标题完整性（任务1 之后至保存）: 通过",
        "\n--- 任务5 验证（条件1–4）---",
        margin_line,
        f"\n  逐张核对 Word 再加载后 extent: {'通过' if task5_ok else '未通过'}",
    )
    if task5_issues:
        print("\n  问题明细:")
        for line in task5_issues:
            print(f"    · {line}")
    elif shape_reports:
        n_32 = sum(1 for r in shape_reports if r["section_3_2_files"])
        print(
            f"\n  摘要: 共 {len(shape_reports)} 张；其中 3.2 Files 区间 {n_32} 张；"
            f" 非 Files 区间 {len(shape_reports) - n_32} 张。"
            "\n  条件1 版心不越界、条件4 等比例：由 extent 与宽高比校验保证。"
            "\n  条件2/3 目标宽度：由算法优先版心后取 min(目标宽,版心宽) 再压高度实现。"
        )


if __name__ == "__main__":
    main()
