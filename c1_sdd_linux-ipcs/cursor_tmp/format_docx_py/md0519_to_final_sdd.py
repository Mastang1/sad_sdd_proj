# -*- coding: utf-8 -*-
r"""
将 ``md_sdd_0519.md`` 全文转换并写入 ``final_sdd.docx``（TF）。

流程：

1. 解析 Markdown：将 ``md_sdd_0519_media/imageN.png`` 按上文 ``### 3.2.x / 3.3.x / 3.4.x`` 映射到
   ``files_32_svgs/`` 或 ``flow_svgs/`` 中的 SVG（缺失则保留原路径并告警）。
2. Mermaid 预处理；**SVG 直嵌 Word**（禁止栅格化为 PNG，见 .cursorrules「插图 SVG 专规」）。
3. Pandoc 生成中间态正文 DOCX（``--reference-doc`` = ``format_refer/format_refer.docx``）。
4. **套用模板**：复制 ``format_refer.docx`` 保留封面，追加 Pandoc 正文，同步页眉/页脚与分节属性
   （``apply_format_refer.merge_body_into_template``）。
5. **HTML 函数设计表**：Pandoc 无法将 ``<table>`` 转为 Word 表格；先从 MD 剥离 HTML 表，
   转换后再用 ``html_table_utils`` 按 colspan/rowspan 插入 Word 表格（与 task-3 规则一致）。
6. 调用 ``format_final_sdd`` 应用 .cursorrules 中的 TF 版式（字体、表格边框、插图宽度等）。
7. HTML 函数表插入与表格适应窗口后，对 ACTIVITY 流程图按 ``scale_flow_diagram_typography``
   逐图设宽（Word 内约五号正文）。
8. 终检页眉/页脚与模板一致；可选跑 ``validate_md_docx_consistency.py``。

**依赖**::

    pip install python-docx docxcompose cairosvg beautifulsoup4 lxml
    pandoc 3.x（PATH 可用）

**使用方式**（仓库根目录 ``c1_sdd_linux-ipcs``）::

    python cursor_tmp/format_docx_py/md0519_to_final_sdd.py
    python cursor_tmp/format_docx_py/md0519_to_final_sdd.py "D:\\path\\final_sdd.docx"

若 ``final_sdd.docx`` 被 Word 占用，将写入 ``final_sdd.generated.docx``。
"""

from __future__ import annotations

import importlib.util
import re
import shutil
import sys
from pathlib import Path

from docx import Document

_SCRIPT_DIR = Path(__file__).resolve().parent
_CURSOR_TMP = _SCRIPT_DIR.parent
if str(_CURSOR_TMP) not in sys.path:
    sys.path.insert(0, str(_CURSOR_TMP))
if str(_SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPT_DIR))
from workspace_paths import (
    WORKSPACE_ROOT,
    MD_SDD_0519,
    FINAL_SDD_DOCX,
    FORMAT_REFER_DOCX,
    PANDOC_MD0519,
    BODY_MD0519,
    FILES_32_SVGS,
    FLOW_SVGS,
    MEDIA_DIR,
    SCRIPTS,
    FORMAT_DOCX_PY,
    pandoc_resource_path_str,
    rel_to_workspace,
)
from apply_format_refer import (
    apply_template_section_properties,
    assert_word_compatible_docx,
    merge_body_into_template,
    sync_header_footer_parts,
    verify_template_headers_footers,
)
from html_table_utils import (
    ensure_sect_pr,
    extract_and_strip_html_tables,
    insert_html_function_tables,
    reapply_all_table_styles,
    apply_table_autofit_window_to_docx,
)
from scale_flow_diagram_typography import (
    REPORT_PATH as FLOW_TYPOGRAPHY_REPORT,
    apply_flow_typography_scale,
    write_report as write_flow_typography_report,
)

WORKSPACE = WORKSPACE_ROOT
SOURCE_MD = MD_SDD_0519
TARGET_DOCX = FINAL_SDD_DOCX
PANDOC_SRC = PANDOC_MD0519
BODY_DOCX = BODY_MD0519
BUILD_SCRIPT = SCRIPTS / "build_final_sdd_docx.py"
FORMAT_SCRIPT = FORMAT_DOCX_PY / "format_final_sdd.py"

_HEADING_RE = re.compile(r"^###\s+(3\.2\.\d+|3\.3\.\d+|3\.4\.\d+)\s+", re.MULTILINE)
_MEDIA_IMG_RE = re.compile(
    r"!\[([^\]]*)\]\(md_sdd_0519_media/image(\d+)\.png\)"
)
_IMG_MD = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def _count_md_images(md_text: str) -> int:
    n = 0
    for _alt, path in _IMG_MD.findall(md_text):
        if not path.startswith(("http://", "https://", "data:")):
            n += 1
    return n


def _assert_md_svg_only(md_text: str) -> None:
    bad = [
        path
        for _alt, path in _IMG_MD.findall(md_text)
        if not path.startswith(("http://", "https://", "data:"))
        and not path.lower().endswith(".svg")
    ]
    if bad:
        sys.exit(
            f"[ERROR] MD 插图必须为 .svg，发现 {len(bad)} 个非 SVG 引用，"
            f"例如: {bad[:5]}"
        )


def _verify_docx_images(docx_path: Path, expected: int) -> int:
    """校验 DOCX 内联插图数量与 display extent；不足或过小则失败退出。"""
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))
    blips = len(doc.element.body.findall(".//" + qn("a:blip")))
    if expected and blips < expected:
        sys.exit(
            f"[ERROR] DOCX 插图未完整嵌入：blips={blips}，MD 引用={expected}。"
            " 请确认 Pandoc --resource-path 含 cursor_tmp 且未错误栅格化为不可达 PNG。"
        )
    min_cx = None
    for inline_el in doc.element.body.findall(".//" + qn("wp:inline")):
        ext_el = inline_el.find(qn("wp:extent"))
        if ext_el is None:
            continue
        cx = int(ext_el.get("cx") or 0)
        min_cx = cx if min_cx is None else min(min_cx, cx)
    if expected and (min_cx is None or min_cx < 100_000):
        sys.exit(
            f"[ERROR] DOCX 插图 display extent 异常（min cx={min_cx} EMU）。"
            " 多为 merge 后缺 w:sectPr 且 format_final_sdd 将图缩成 1×1；"
            " 请确认 ensure_sect_pr 已 save，并重跑转换。"
        )
    return blips


def _collect_md_svg_aspect_ratios(md_text: str) -> list[float]:
    """按 MD 引用顺序，从磁盘 SVG viewBox 收集 cy/cx。"""
    from svg_extent_utils import parse_svg_viewbox_wh

    ratios: list[float] = []
    for _alt, path in _IMG_MD.findall(md_text):
        if path.startswith(("http://", "https://", "data:")):
            continue
        svg_path = (WORKSPACE / path).resolve()
        if not svg_path.is_file():
            continue
        wh = parse_svg_viewbox_wh(svg_path.read_text(encoding="utf-8", errors="ignore"))
        if wh is None:
            continue
        w, h = wh
        ratios.append(h / w)
    return ratios


def _collect_inline_aspect_ratios(docx_path: Path) -> list[float]:
    """按文档顺序收集各内联图 cy/cx（wp:extent）。"""
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))
    ratios: list[float] = []
    for inline_el in doc.element.body.findall(".//" + qn("wp:inline")):
        ext_el = inline_el.find(qn("wp:extent"))
        if ext_el is None:
            continue
        cx = int(ext_el.get("cx") or 0)
        cy = int(ext_el.get("cy") or 0)
        if cx > 0 and cy > 0:
            ratios.append(cy / cx)
    return ratios


def _verify_aspect_ratios_preserved(
    expected: list[float], actual: list[float], tolerance: float = 0.02
) -> None:
    """终检：DOCX 插图 cy/cx 须与 MD 源 SVG viewBox 一致（§5.7/§6.7 时序图等）。"""
    if len(expected) != len(actual):
        sys.exit(
            f"[ERROR] 宽高比校验：插图数量不一致 MD={len(expected)} DOCX={len(actual)}。"
        )
    bad: list[tuple[int, float, float]] = []
    for i, (exp, act) in enumerate(zip(expected, actual)):
        if exp <= 0:
            continue
        if abs(act - exp) / exp > tolerance:
            bad.append((i + 1, exp, act))
    if bad:
        ex = bad[0]
        sys.exit(
            f"[ERROR] {len(bad)} 张图宽高比与源 SVG viewBox 不一致。"
            f" 例：图 #{ex[0]} viewBox={ex[1]:.4f} docx={ex[2]:.4f}。"
            " format_final_sdd 须从 word/media 内嵌 SVG viewBox 计算 extent。"
        )


def _load_build():
    spec = importlib.util.spec_from_file_location("bfd", BUILD_SCRIPT)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"cannot load {BUILD_SCRIPT}")
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def _svg_for_section(section: str) -> Path | None:
    """3.2.2 -> files_32_svgs/3_2_2.svg ; 3.3.1 / 3.4.23 -> flow_svgs/3_3_1.svg"""
    parts = section.split(".")
    if len(parts) != 3:
        return None
    major, minor, patch = parts
    if major == "3" and minor == "2":
        p = FILES_32_SVGS / f"3_2_{patch}.svg"
    elif major == "3" and minor in ("3", "4"):
        p = FLOW_SVGS / f"3_{minor}_{patch}.svg"
        if not p.is_file() and minor == "4":
            alt = FLOW_SVGS / f"tx_3_4_{patch}.svg"
            if alt.is_file():
                p = alt
    else:
        return None
    return p if p.is_file() else None


def resolve_media_images(md_text: str) -> str:
    """按最近 ### 3.x.y 标题，将 md_sdd_0519_media 占位图改为本地 SVG 路径。"""
    last_section: str | None = None
    out_lines: list[str] = []
    missing: list[str] = []

    for line in md_text.splitlines(keepends=True):
        hm = _HEADING_RE.match(line)
        if hm:
            last_section = hm.group(1)

        def repl(m: re.Match[str]) -> str:
            nonlocal last_section
            alt = m.group(1)
            num = m.group(2)
            if last_section:
                svg = _svg_for_section(last_section)
                if svg:
                    rel = rel_to_workspace(svg).as_posix()
                    return f"![{alt}]({rel})"
            missing.append(f"image{num}.png (section={last_section})")
            return m.group(0)

        out_lines.append(_MEDIA_IMG_RE.sub(repl, line))

    if missing:
        print(f"[warn] {len(missing)} image(s) could not be mapped to SVG, e.g. {missing[:5]}")
    return "".join(out_lines)


def ensure_media_from_docx(docx_path: Path) -> None:
    """从现有 TF 解出 word/media/* 到 md_sdd_0519_media/，供未映射的 imageN.png 使用。"""
    import zipfile

    media_dir = MEDIA_DIR
    media_dir.mkdir(parents=True, exist_ok=True)
    if not docx_path.is_file():
        return
    with zipfile.ZipFile(docx_path) as z:
        for name in z.namelist():
            if not name.startswith("word/media/"):
                continue
            base = Path(name).name
            target = media_dir / base
            if not target.is_file():
                target.write_bytes(z.read(name))


def _special_image_fixes(md_text: str) -> str:
    """4.2.1 平台头文件等仍引用 image40；映射到 3_2_9.svg。"""
    p = FILES_32_SVGS / "3_2_9.svg"
    if p.is_file():
        rel = rel_to_workspace(p)
        md_text = md_text.replace("md_sdd_0519_media/image40.png", rel)
    return md_text


def run_pandoc_extended(
    bfd, md_path: Path, out_docx: Path, reference_docx: Path | None
) -> None:
    """Pandoc 转 DOCX；插图保持 SVG 直嵌（禁止先栅格化为 PNG）。"""
    rpath = pandoc_resource_path_str()
    import subprocess

    cmd = [
        "pandoc",
        str(md_path),
        "-o",
        str(out_docx),
        "-f",
        "markdown+pipe_tables+raw_html+smart+yaml_metadata_block",
        "-t",
        "docx",
        "--resource-path",
        rpath,
        "--standalone",
    ]
    if reference_docx and reference_docx.is_file():
        cmd.extend(["--reference-doc", str(reference_docx)])
        print(f"[info] Pandoc reference-doc: {reference_docx.name}")
    subprocess.run(cmd, cwd=str(WORKSPACE), check=True)
    print(f"[info] Pandoc → {out_docx} (SVG direct embed, resource-path OK)")


def run_format_final_sdd(docx_path: Path) -> None:
    spec = importlib.util.spec_from_file_location("format_final_sdd", FORMAT_SCRIPT)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"cannot load {FORMAT_SCRIPT}")
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    old_argv = sys.argv[:]
    try:
        sys.argv = [str(FORMAT_SCRIPT), str(docx_path)]
        mod.main()
    finally:
        sys.argv = old_argv


def main() -> None:
    out_docx = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else TARGET_DOCX.resolve()
    if not SOURCE_MD.is_file():
        sys.exit(f"missing: {SOURCE_MD}")
    if not FORMAT_REFER_DOCX.is_file():
        sys.exit(f"missing format template: {FORMAT_REFER_DOCX}")

    bfd = _load_build()
    bfd.ROOT = WORKSPACE
    bfd.TARGET_DOCX = out_docx
    bfd.SOURCE_MD = SOURCE_MD

    ensure_media_from_docx(FORMAT_REFER_DOCX)

    md = SOURCE_MD.read_text(encoding="utf-8")
    md = resolve_media_images(md)
    md = _special_image_fixes(md)
    _assert_md_svg_only(md)
    expected_images = _count_md_images(md)
    expected_aspects = _collect_md_svg_aspect_ratios(md)
    print(f"[info] MD image refs (SVG): {expected_images}")
    md, html_table_blocks = extract_and_strip_html_tables(md)
    if html_table_blocks:
        print(f"[info] Stripped {len(html_table_blocks)} HTML function table(s) for post-insert")
    md = bfd.preprocess_mermaid(md)
    # SVG 直嵌：勿调用 rasterize_svg_refs_for_docx（Windows 无 rsvg 时 PNG 路径错误会导致插图丢失）
    PANDOC_SRC.write_text(md, encoding="utf-8")

    run_pandoc_extended(bfd, PANDOC_SRC, BODY_DOCX, FORMAT_REFER_DOCX)
    _verify_docx_images(BODY_DOCX, expected_images)

    merged_path = merge_body_into_template(
        FORMAT_REFER_DOCX, BODY_DOCX, out_docx, workspace_root=WORKSPACE
    )

    alt = out_docx.with_name(out_docx.stem + ".generated.docx")
    if merged_path != out_docx.resolve() and alt.is_file():
        print(f"[info] Using merge output {merged_path}")

    if merged_path.is_file():
        _merged_doc = Document(str(merged_path))
        if ensure_sect_pr(_merged_doc):
            _merged_doc.save(str(merged_path))
            sync_header_footer_parts(merged_path, FORMAT_REFER_DOCX)
            apply_template_section_properties(merged_path, FORMAT_REFER_DOCX)
            sync_header_footer_parts(merged_path, FORMAT_REFER_DOCX)
            print(f"[info] Re-applied template sectPr after ensure_sect_pr → {merged_path.name}")
    saved = merged_path

    for p in (PANDOC_SRC, BODY_DOCX):
        if p and p.is_file():
            p.unlink()

    print("[info] Applying format_final_sdd (TF styles)...")
    run_format_final_sdd(saved)

    if html_table_blocks:
        inserted, missing = insert_html_function_tables(saved, html_table_blocks)
        print(f"[info] Inserted {inserted} HTML function Word table(s)")
        if missing:
            print(f"[warn] Could not locate Heading 3 for: {missing[:8]}...")
        reapply_all_table_styles(saved)

    n_autofit = apply_table_autofit_window_to_docx(saved)
    print(f"[info] Table autofit window: {n_autofit} table(s)")

    scaled, skipped, flow_reports = apply_flow_typography_scale(saved)
    write_flow_typography_report(FLOW_TYPOGRAPHY_REPORT, saved, scaled, skipped, flow_reports)
    print(
        f"[info] ACTIVITY flow typography (五号): {scaled} diagram(s); "
        f"skipped {skipped} inline image(s)"
    )
    print(f"[info] Flow typography report: {FLOW_TYPOGRAPHY_REPORT}")

    aspects_after = _collect_inline_aspect_ratios(saved)
    _verify_aspect_ratios_preserved(expected_aspects, aspects_after)
    print(f"[info] Aspect ratios match SVG viewBox for {len(expected_aspects)} image(s)")

    sync_header_footer_parts(saved, FORMAT_REFER_DOCX)
    apply_template_section_properties(saved, FORMAT_REFER_DOCX)
    sync_header_footer_parts(saved, FORMAT_REFER_DOCX)
    verify_template_headers_footers(saved, FORMAT_REFER_DOCX)
    assert_word_compatible_docx(saved)
    final_blips = _verify_docx_images(saved, expected_images)
    print(f"[info] DOCX inline images verified: {final_blips}/{expected_images}")

    if saved != out_docx.resolve():
        try:
            shutil.copy2(saved, out_docx)
            print(f"[info] Replaced locked target → {out_docx}")
            saved = out_docx.resolve()
        except OSError as e:
            print(
                f"[ERROR] 目标 {out_docx} 仍被占用，当前有效输出为 {saved}。"
                f" 请关闭 Word 后执行: copy /Y \"{saved}\" \"{out_docx}\"",
                file=sys.stderr,
            )

    print(f"[done] {saved}")
    print("[hint] Run: python cursor_tmp/format_docx_py/validate_md_docx_consistency.py")
    if saved != out_docx and out_docx.is_file():
        print(f"[hint] Close Word and replace {out_docx} with {saved}")


if __name__ == "__main__":
    main()
