# -*- coding: utf-8 -*-
"""
将 Pandoc 中间态正文合并进 ``format_refer/format_refer.docx`` 模板，
并同步页眉/页脚与分节属性，使交付 DOCX 与模板完全一致。

**关键约束**：禁止对 ``word/document.xml`` / ``document.xml.rels`` 使用
``xml.etree.ElementTree.tostring()`` 写回（会把 ``w:`` 前缀变成 ``ns0:``，
Word 无法打开）。sectPr 注入须走 ``python-docx``；HF 部件只做 ZIP 原始字节覆盖。
"""

from __future__ import annotations

import copy
import re
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_CURSOR_TMP = Path(__file__).resolve().parents[1]
if str(_CURSOR_TMP) not in sys.path:
    sys.path.insert(0, str(_CURSOR_TMP))
from workspace_paths import FORMAT_REFER_DOCX, WORKSPACE_ROOT

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}
HF_REL_TYPES = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header",
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer",
)
_RELationship_RE = re.compile(r"<Relationship\s+[^>]+/>")


def _load_build():
    import importlib.util

    build_script = _CURSOR_TMP / "scripts" / "build_final_sdd_docx.py"
    spec = importlib.util.spec_from_file_location("bfd", build_script)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"cannot load {build_script}")
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def assert_word_compatible_docx(docx_path: Path) -> None:
    """Word 打不开的常见根因：document.xml 被写成 ns0: 前缀。"""
    with zipfile.ZipFile(docx_path) as z:
        head = z.read("word/document.xml")[:4096]
    if b"ns0:" in head or b"ns1:" in head:
        sys.exit(
            "[ERROR] word/document.xml 含 ns0/ns1 前缀，Word 无法打开。"
            " 请使用 python-docx 写 sectPr，勿用 ElementTree.tostring。"
        )
    if not head.lstrip().startswith(b"<?xml"):
        sys.exit("[ERROR] word/document.xml 缺少 XML 声明")


def _first_page_break_paragraph_el(body) -> OxmlElement | None:
    for child in body:
        if child.tag != qn("w:p"):
            continue
        for br in child.iter(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return child
    return None


def _extract_sect_pr_elements(template_path: Path) -> tuple[OxmlElement, OxmlElement]:
    """从模板提取 cover/body 两套 sectPr（python-docx Oxml 深拷贝）。"""
    tpl = Document(str(template_path))
    body = tpl.element.body

    cover_sp: OxmlElement | None = None
    for child in body:
        if child.tag != qn("w:p"):
            continue
        p_pr = child.find(qn("w:pPr"))
        if p_pr is None:
            continue
        sp = p_pr.find(qn("w:sectPr"))
        if sp is not None:
            cover_sp = copy.deepcopy(sp)
            break

    body_sp_el = body.find(qn("w:sectPr"))
    if cover_sp is None or body_sp_el is None:
        raise RuntimeError(
            f"{template_path} 缺少 cover/body 分节属性（inline sectPr 或 body sectPr）"
        )
    return cover_sp, copy.deepcopy(body_sp_el)


def _ensure_p_pr_el(p_el: OxmlElement) -> OxmlElement:
    p_pr = p_el.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        p_el.insert(0, p_pr)
    return p_pr


def apply_template_section_properties(
    docx_path: Path, template_path: Path = FORMAT_REFER_DOCX
) -> None:
    """将模板 cover/body 分节属性写回目标 DOCX（python-docx 保存，保留 w: 前缀）。"""
    cover_sp, body_sp = _extract_sect_pr_elements(template_path)
    doc = Document(str(docx_path))
    body = doc.element.body

    cover_break_p = _first_page_break_paragraph_el(body)
    if cover_break_p is not None:
        p_pr = _ensure_p_pr_el(cover_break_p)
        old = p_pr.find(qn("w:sectPr"))
        if old is not None:
            p_pr.remove(old)
        p_pr.append(copy.deepcopy(cover_sp))
    else:
        print(
            "[warn] 未发现封面分页符，封面节 sectPr 未注入；"
            "请确认模板含手动分页符。",
            file=sys.stderr,
        )

    stripped = 0
    for child in body:
        if child.tag != qn("w:p") or child is cover_break_p:
            continue
        p_pr = child.find(qn("w:pPr"))
        if p_pr is None:
            continue
        sp = p_pr.find(qn("w:sectPr"))
        if sp is not None:
            p_pr.remove(sp)
            stripped += 1

    old_body_sp = body.find(qn("w:sectPr"))
    if old_body_sp is not None:
        body.remove(old_body_sp)
    body.append(copy.deepcopy(body_sp))

    doc.save(str(docx_path))
    assert_word_compatible_docx(docx_path)
    if stripped:
        print(f"[info] Removed {stripped} inline sectPr from appended body")


def _replace_zip_parts(docx_path: Path, replacements: dict[str, bytes]) -> None:
    """替换 DOCX 内若干 part（保留其余 entry 的 ZipInfo）。"""
    if not replacements:
        return
    fd, tmp_name = tempfile.mkstemp(suffix=".docx", dir=str(docx_path.parent))
    import os

    os.close(fd)
    tmp_path = Path(tmp_name)
    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            written = set()
            for item in zin.infolist():
                if item.filename in replacements:
                    zout.writestr(item, replacements[item.filename])
                    written.add(item.filename)
                else:
                    zout.writestr(item, zin.read(item.filename))
            for name, data in replacements.items():
                if name not in written:
                    zout.writestr(name, data)
        shutil.copy2(tmp_path, docx_path)
    finally:
        tmp_path.unlink(missing_ok=True)


def _hf_part_names(zip_names: set[str]) -> set[str]:
    out: set[str] = set()
    for name in zip_names:
        base = Path(name).name
        if base.startswith("header") and base.endswith(".xml"):
            out.add(name)
        elif base.startswith("footer") and base.endswith(".xml"):
            out.add(name)
        elif name.startswith("word/_rels/header") and name.endswith(".xml.rels"):
            out.add(name)
        elif name.startswith("word/_rels/footer") and name.endswith(".xml.rels"):
            out.add(name)
    return out


def _resolve_part_target(rels_path: str, target: str) -> str:
    base_dir = Path(rels_path).parent.parent
    parts: list[str] = []
    for seg in (base_dir / target).as_posix().split("/"):
        if seg == "..":
            if parts:
                parts.pop()
        elif seg and seg != ".":
            parts.append(seg)
    return "/".join(parts)


def _collect_hf_media_from_rels(
    z: zipfile.ZipFile, rels_path: str, collected: set[str]
) -> None:
    if rels_path not in z.namelist():
        return
    rels = ET.fromstring(z.read(rels_path))
    for rel in rels:
        target = rel.get("Target", "")
        if not target or target.startswith("#"):
            continue
        resolved = _resolve_part_target(rels_path, target)
        if resolved in z.namelist() and resolved not in collected:
            collected.add(resolved)
            if resolved.startswith("word/_rels/") and resolved.endswith(".rels"):
                _collect_hf_media_from_rels(z, resolved, collected)


def _merge_document_hf_rels(doc_rels: bytes, tpl_rels: bytes) -> bytes:
    """追加模板中缺失的 header/footer Relationship（字符串级，不改命名空间前缀）。"""
    doc_text = doc_rels.decode("utf-8")
    tpl_text = tpl_rels.decode("utf-8")
    existing_targets = set(re.findall(r'Target="([^"]+)"', doc_text))
    existing_ids = set(re.findall(r'Id="([^"]+)"', doc_text))
    to_add: list[str] = []
    for rel in _RELationship_RE.findall(tpl_text):
        if not any(t in rel for t in ("/header", "/footer")):
            continue
        m_t = re.search(r'Target="([^"]+)"', rel)
        m_i = re.search(r'Id="([^"]+)"', rel)
        if m_t and m_t.group(1) in existing_targets:
            continue
        if m_i and m_i.group(1) in existing_ids:
            continue
        to_add.append(rel)
    if not to_add:
        return doc_rels
    pos = doc_text.rfind("</Relationships>")
    if pos < 0:
        return doc_rels
    merged = doc_text[:pos] + "".join(to_add) + doc_text[pos:]
    return merged.encode("utf-8")


def sync_header_footer_parts(
    docx_path: Path, template_path: Path = FORMAT_REFER_DOCX
) -> None:
    """从模板复制 header/footer 及其 rels、媒体到目标 DOCX（原始字节覆盖）。"""
    rels_name = "word/_rels/document.xml.rels"
    with zipfile.ZipFile(template_path, "r") as ztpl, zipfile.ZipFile(
        docx_path, "r"
    ) as zdoc:
        tpl_names = set(ztpl.namelist())
        doc_names = set(zdoc.namelist())
        hf_parts = _hf_part_names(tpl_names)
        for part in list(hf_parts):
            if part.startswith("word/header") or part.startswith("word/footer"):
                rels = "word/_rels/" + Path(part).name + ".rels"
                if rels in tpl_names:
                    hf_parts.add(rels)
        extra: set[str] = set()
        for part in sorted(hf_parts):
            if part.startswith("word/_rels/"):
                _collect_hf_media_from_rels(ztpl, part, extra)
        hf_parts |= extra

        replacements: dict[str, bytes] = {}
        for part in sorted(hf_parts):
            if part not in tpl_names:
                continue
            data = ztpl.read(part)
            if part not in doc_names or zdoc.read(part) != data:
                replacements[part] = data

        new_rels = _merge_document_hf_rels(
            zdoc.read(rels_name), ztpl.read(rels_name)
        )
        if new_rels != zdoc.read(rels_name):
            replacements[rels_name] = new_rels

    if not replacements:
        return
    _replace_zip_parts(docx_path, replacements)
    print(f"[info] Synced {len(replacements)} header/footer part(s) from template")


def _sect_pr_signature_from_el(sect_pr) -> tuple:
    if sect_pr is None:
        return ()
    refs: list[tuple[str, str, str]] = []
    for child in sect_pr:
        tag = child.tag.split("}")[-1]
        if tag.endswith("Reference"):
            refs.append((tag, child.get(qn("w:type"), ""), child.get(qn("w:id"), "")))
        elif tag == "titlePg":
            refs.append(("titlePg", "", ""))
    return tuple(sorted(refs))


def _extract_sect_pr_for_verify(template_path: Path) -> tuple:
    cover_sp, body_sp = _extract_sect_pr_elements(template_path)
    return _sect_pr_signature_from_el(cover_sp), _sect_pr_signature_from_el(body_sp)


def verify_template_headers_footers(
    docx_path: Path, template_path: Path = FORMAT_REFER_DOCX
) -> None:
    """校验 HF 部件与分节引用是否与模板一致。"""
    cover_sig_tpl, body_sig_tpl = _extract_sect_pr_for_verify(template_path)

    with zipfile.ZipFile(template_path, "r") as ztpl, zipfile.ZipFile(
        docx_path, "r"
    ) as zdoc:
        tpl_hf = _hf_part_names(set(ztpl.namelist()))
        doc_hf = _hf_part_names(set(zdoc.namelist()))
        missing_parts = tpl_hf - doc_hf
        if missing_parts:
            sys.exit(
                f"[ERROR] DOCX 缺少模板页眉/页脚部件: {sorted(missing_parts)[:6]}"
            )
        for part in tpl_hf:
            if "rels" in part:
                continue
            if ztpl.read(part) != zdoc.read(part):
                sys.exit(f"[ERROR] 页眉/页脚内容与模板不一致: {part}")

    assert_word_compatible_docx(docx_path)
    doc = Document(str(docx_path))
    body = doc.element.body
    cover_break_p = _first_page_break_paragraph_el(body)
    cover_sp = None
    if cover_break_p is not None:
        p_pr = cover_break_p.find(qn("w:pPr"))
        if p_pr is not None:
            cover_sp = p_pr.find(qn("w:sectPr"))
    body_sp = body.find(qn("w:sectPr"))

    if _sect_pr_signature_from_el(cover_sp) != cover_sig_tpl:
        sys.exit("[ERROR] 封面节页眉/页脚引用与模板不一致")
    if _sect_pr_signature_from_el(body_sp) != body_sig_tpl:
        sys.exit("[ERROR] 正文节页眉/页脚引用与模板不一致")

    print("[info] Header/footer verification: PASS (matches template)")


def merge_body_into_template(
    template_path: Path,
    body_path: Path,
    out_path: Path,
    *,
    workspace_root: Path = WORKSPACE_ROOT,
) -> Path:
    """
    复制模板、保留封面、追加 Pandoc 正文，并同步页眉页脚。

    返回实际写入路径（``out_path`` 或 ``*.generated.docx``）。
    """
    if not template_path.is_file():
        sys.exit(f"[ERROR] 模板不存在: {template_path}")
    if not body_path.is_file():
        sys.exit(f"[ERROR] 正文 DOCX 不存在: {body_path}")

    bfd = _load_build()
    bfd.ROOT = workspace_root

    fd, cover_copy = tempfile.mkstemp(suffix=".docx", dir=str(workspace_root))
    import os

    os.close(fd)
    cover_copy_p = Path(cover_copy)
    tmp = Path(tempfile.mkstemp(suffix=".docx", dir=str(workspace_root))[1])
    try:
        shutil.copy2(template_path, cover_copy_p)
        master = Document(str(cover_copy_p))
        bfd.trim_keeps_cover_removes_rest(master)
        from docxcompose.composer import Composer

        composer = Composer(master)
        composer.append(Document(str(body_path)))
        composer.save(str(tmp))
        del composer, master

        try:
            shutil.copyfile(str(tmp), str(out_path))
            print(f"[info] Merged template cover + Pandoc body → {out_path}")
            saved = out_path
        except OSError as e:
            alt = out_path.with_name(out_path.stem + ".generated.docx")
            shutil.copyfile(str(tmp), str(alt))
            print(
                f"[warn] Cannot overwrite {out_path} ({e}); wrote {alt}.",
                file=sys.stderr,
            )
            saved = alt
    finally:
        bfd._unlink_retry(tmp)
        bfd._unlink_retry(cover_copy_p)

    assert_word_compatible_docx(saved)
    sync_header_footer_parts(saved, template_path)
    apply_template_section_properties(saved, template_path)
    sync_header_footer_parts(saved, template_path)
    verify_template_headers_footers(saved, template_path)
    return saved


def main() -> None:
    import argparse

    from workspace_paths import FINAL_SDD_DOCX

    ap = argparse.ArgumentParser(description="Merge Pandoc body into format_refer template")
    ap.add_argument("body", type=Path, help="Pandoc 中间态 DOCX")
    ap.add_argument("-o", "--output", type=Path, default=FINAL_SDD_DOCX)
    ap.add_argument("-t", "--template", type=Path, default=FORMAT_REFER_DOCX)
    args = ap.parse_args()
    merge_body_into_template(
        args.template.resolve(), args.body.resolve(), args.output.resolve()
    )


if __name__ == "__main__":
    main()
