# -*- coding: utf-8 -*-
"""HTML 函数设计表 → Word 表格（含 colspan/rowspan）共用工具。"""

from __future__ import annotations

import re
from typing import Iterable

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table

from format_final_sdd import (
    _EAST_ASIA_FONT,
    _TABLE_FONT_PT,
    _WESTERN_FONT,
    apply_solid_table_borders,
    _set_run_font,
)

_HTML_FUNC_TABLE = re.compile(
    r"(###\s+\d+\.\d+(?:\.\d+)?\s+\S+)\s*\n\n"
    r'(<table border="1"[^>]*>.*?</table>)\s*',
    re.DOTALL,
)


def parse_html_placements(table_html: str) -> tuple[int, int, list[tuple[int, int, int, int, str]]]:
    soup = BeautifulSoup(table_html, "html.parser")
    table = soup.find("table")
    if not table:
        return 0, 0, []
    occupied: dict[tuple[int, int], bool] = {}
    placements: list[tuple[int, int, int, int, str]] = []
    cur_r = 0
    for tr in table.find_all("tr"):
        tds = tr.find_all("td")
        if not tds:
            continue
        cur_c = 0
        for td in tds:
            while (cur_r, cur_c) in occupied:
                cur_c += 1
            rs = int(td.get("rowspan", 1))
            cs = int(td.get("colspan", 1))
            text = td.get_text(separator="\n").strip()
            placements.append((cur_r, cur_c, rs, cs, text))
            for dr in range(rs):
                for dc in range(cs):
                    occupied[(cur_r + dr, cur_c + dc)] = True
            cur_c += cs
        cur_r += 1
    nrows = cur_r
    ncols = max((c + cs for _, c, _, cs, _ in placements), default=0)
    return nrows, ncols, placements


def extract_and_strip_html_tables(md_text: str) -> tuple[str, list[tuple[str, str, str]]]:
    """返回 (新 MD 文本, [(func_name, heading_line, table_html), ...])。"""
    blocks: list[tuple[str, str, str]] = []

    def repl(m: re.Match[str]) -> str:
        heading = m.group(1).strip()
        table_html = m.group(2)
        # Match by full H3 text (e.g. "6.3.3 ipcsOsInit"), not bare func name — supports duplicates.
        h3_key = re.sub(r"^###\s+", "", heading).strip()
        blocks.append((h3_key, heading, table_html))
        return heading + "\n\n"

    new_md = _HTML_FUNC_TABLE.sub(repl, md_text)
    return new_md, blocks


def _paragraph_style_name(doc: Document, p_el) -> str | None:
    ppr = p_el.pPr
    if ppr is None or ppr.pStyle is None:
        return None
    try:
        from docx.enum.style import WD_STYLE_TYPE

        return doc.styles.get_by_id(ppr.pStyle.val, WD_STYLE_TYPE.PARAGRAPH).name
    except (KeyError, ValueError, AttributeError):
        return None


def find_h3_elements_in_order(doc: Document) -> list[tuple[object, str]]:
    items: list[tuple[object, str]] = []
    for p in doc.paragraphs:
        if _paragraph_style_name(doc, p._element) != "Heading 3":
            continue
        items.append((p._element, p.text.strip()))
    return items


def ensure_sect_pr(doc: Document) -> bool:
    """docxcompose 合并后可能缺少 w:sectPr，导致 add_table / 版心计算失败。返回是否新插入 sectPr。"""
    body = doc.element.body
    if body.findall(qn("w:sectPr")):
        return False
    sect_pr = OxmlElement("w:sectPr")
    pg_sz = OxmlElement("w:pgSz")
    pg_sz.set(qn("w:w"), "11906")
    pg_sz.set(qn("w:h"), "16838")
    sect_pr.append(pg_sz)
    pg_mar = OxmlElement("w:pgMar")
    for edge, twips in (
        ("top", "1440"),
        ("right", "1440"),
        ("bottom", "1440"),
        ("left", "1440"),
    ):
        pg_mar.set(qn(f"w:{edge}"), twips)
    sect_pr.append(pg_mar)
    body.append(sect_pr)
    return True


def insert_table_after_element(
    doc: Document,
    anchor_el,
    nrows: int,
    ncols: int,
    placements: Iterable[tuple[int, int, int, int, str]],
):
    temp = doc.add_table(rows=nrows, cols=ncols)
    tbl_el = temp._tbl
    temp._tbl.getparent().remove(tbl_el)
    anchor_el.addnext(tbl_el)
    table = Table(tbl_el, doc)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r, c, rs, cs, text in placements:
        cell = table.cell(r, c)
        if rs > 1 or cs > 1:
            cell.merge(table.cell(r + rs - 1, c + cs - 1))
        cell.text = text or " "
        for para in cell.paragraphs:
            for run in para.runs:
                _set_run_font(
                    run,
                    western=_WESTERN_FONT,
                    east_asia=_EAST_ASIA_FONT,
                    size_pt=_TABLE_FONT_PT,
                    bold=False,
                )
    apply_solid_table_borders(table)
    return tbl_el


def insert_html_function_tables(
    docx_path, blocks: list[tuple[str, str, str]]
) -> tuple[int, list[str]]:
    """按文档中 Heading 3 顺序插入 HTML 函数表（支持同名函数多次出现）。"""
    doc = Document(str(docx_path))
    ensure_sect_pr(doc)
    h3_list = find_h3_elements_in_order(doc)
    pos = 0
    inserted = 0
    missing: list[str] = []

    def _norm_name(s: str) -> str:
        return re.sub(r"\s+", "", s).lower()

    for h3_key, heading, table_html in blocks:
        func_name = h3_key.split()[-1]
        found = None
        for i in range(pos, len(h3_list)):
            t = h3_list[i][1]
            if _norm_name(t) not in (_norm_name(h3_key), _norm_name(func_name)):
                continue
            h_el = h3_list[i][0]
            nxt = h_el.getnext()
            if nxt is not None and nxt.tag == qn("w:tbl"):
                continue
            found = i
            break
        if found is None:
            missing.append(heading)
            continue
        h_el = h3_list[found][0]
        pos = found + 1
        nrows, ncols, placements = parse_html_placements(table_html)
        if nrows == 0:
            missing.append(heading)
            continue
        insert_table_after_element(doc, h_el, nrows, ncols, placements)
        inserted += 1
    doc.save(str(docx_path))
    return inserted, missing


def reapply_all_table_styles(docx_path) -> None:
    from format_final_sdd import apply_all_table_borders, format_all_table_cells

    doc = Document(str(docx_path))
    apply_all_table_borders(doc)
    format_all_table_cells(doc)
    doc.save(str(docx_path))


def apply_table_autofit_window_to_docx(docx_path) -> int:
    """仅对已有 docx 全部表格设置「适应窗口」，不改动边框/字体/正文。"""
    from format_final_sdd import apply_all_table_autofit_window

    doc = Document(str(docx_path))
    n = apply_all_table_autofit_window(doc)
    doc.save(str(docx_path))
    return n
