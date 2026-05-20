# -*- coding: utf-8 -*-
r"""
仅执行任务 6（TF / final_sdd.docx）
==================================

依照 ``prompt-format.md`` **任务 6**：

- 遍历**插图及图表**所在段落，将其在文档中的布局设为：**行首无缩进**、**居中对齐**（相对页面，
  **不是**表格单元格内对齐）。
- 本实现与 ``format_final_sdd.format_figure_chart_paragraphs_layout`` 一致：仅处理
  ``document.paragraphs``（主文档正文流），**不**修改单元格内段落，符合「不是单元格内容」的约定。
- **其他任务一律不执行**（满足「只能修改任务指定内容」）。

**依赖**::

    pip install python-docx

**使用方式**（仓库根目录 ``c1_sdd``）::

    python format_docx_py/format_tf_task_6.py
    python format_docx_py/format_tf_task_6.py "C:\\tangyapeng\\docs\\StarGather\\c1_sdd\\final_sdd.docx"

脚本会写入 TF 并**再次打开**文件，调用 ``verify_task6_figure_paragraph_layout`` 输出核对结果。
若 Word 正打开该文件，请先关闭。

版本：与 TF / prompt-format 配套。
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

from format_final_sdd import (
    DEFAULT_TF,
    format_figure_chart_paragraphs_layout,
    verify_task6_figure_paragraph_layout,
)


def _tf_path() -> Path:
    """
    解析命令行中的 TF 路径；无参数时使用默认 ``final_sdd.docx``。

    :return: 绝对路径
    """
    if len(sys.argv) > 1:
        return Path(sys.argv[1]).resolve()
    return DEFAULT_TF.resolve()


def main() -> None:
    """执行任务 6、保存、再加载并核对。"""
    tf = _tf_path()
    if not tf.is_file():
        print(f"未找到 TF：{tf}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(tf))
    n = format_figure_chart_paragraphs_layout(doc)
    doc.save(str(tf))

    doc2 = Document(str(tf))
    ok, issues, stats = verify_task6_figure_paragraph_layout(doc2)

    print("处理完成并已保存：", tf)
    print(f"  任务6 已设置主文档流中含图段落: {n}")
    print("--- 任务6 需求核对 ---")
    print(
        f"  主文档流含图/表段落数: {stats['main_story_figure_paragraphs']} "
        f"（其中通过校验: {stats['main_story_passed']}）"
    )
    print(
        f"  表格单元格内含图段落数（任务6未改单元格对齐）: "
        f"{stats['table_cell_figure_paragraphs']}"
    )
    print(f"  结论: {'需求已落实（主文档流范围）' if ok else '存在未通过项，见下列明细'}")
    if issues:
        for line in issues:
            print(f"    · {line}")
        sys.exit(1)


if __name__ == "__main__":
    main()
