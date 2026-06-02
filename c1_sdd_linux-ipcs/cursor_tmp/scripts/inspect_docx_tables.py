# -*- coding: utf-8 -*-
import sys
import zipfile
from collections import Counter

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.oxml.ns import qn

z = zipfile.ZipFile("final_sdd.docx")
styles_xml = z.read("word/styles.xml").decode("utf-8")
for sid in ["Compact", "Normal", "TableContents", "Table"]:
    marker = f'w:styleId="{sid}"'
    idx = styles_xml.find(marker)
    if idx >= 0:
        end = styles_xml.find("</w:style>", idx)
        chunk = styles_xml[idx : end + 10]
        print(sid, "numPr=", "numPr" in chunk)

doc = Document("final_sdd.docx")
found = 0
for ti, t in enumerate(doc.tables):
    if len(t.columns) < 5:
        continue
    all_text = "\n".join(c.text for row in t.rows for c in row.cells)
    if "\u51fd\u6570\u539f\u578b" not in all_text:
        continue
    found += 1
    styles = Counter()
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                pPr = p._element.find(qn("w:pPr"))
                sid = (
                    pPr.pStyle.val
                    if (pPr is not None and pPr.pStyle is not None)
                    else "(none)"
                )
                styles[sid] += 1
    print("table", ti, "rows", len(t.rows), "styles", dict(styles))
    print(" row0:", repr(t.rows[0].cells[0].text[:50]))
    if found >= 3:
        break
