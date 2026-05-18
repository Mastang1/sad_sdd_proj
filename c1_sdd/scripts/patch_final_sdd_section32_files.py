# -*- coding: utf-8 -*-
"""Replace ONLY the DOCX span ``Heading 2 : 3.2 Files`` through before ``3.3 External``.

Rebuilds that span from ``ipcs_sdd.md`` (``## 3.2 Files`` … before ``## 3.3``) through
Pandoc + CairoSVG raster of ``files_32_svgs/*.svg``. Other chapters are left as-is.

Uses ``docxcompose.Composer`` so embedded diagram PNGs keep valid package relationships
(in-tree ``deepcopy`` of ``w:p`` would break ``r:embed`` links).

Requires: pandoc, cairosvg, python-docx, docxcompose

Usage::

    cd c:\\tangyapeng\\docs\\StarGather\\c1_sdd
    python scripts\\patch_final_sdd_section32_files.py
"""
from __future__ import annotations

import importlib.util
import re
import shutil
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docxcompose.composer import Composer

ROOT = Path(__file__).resolve().parents[1]
FINAL_DOCX = ROOT / "final_sdd.docx"
SOURCE_MD = ROOT / "ipcs_sdd.md"
TMP_MD = ROOT / "_patch_section32.md"
TMP_DOCX = ROOT / "_patch_section32_generated.docx"
PATCH_HEAD = ROOT / "_patch_head_trim.docx"
PATCH_TAIL = ROOT / "_patch_tail_trim.docx"
WORK_COPY = ROOT / "_patch_master_work.docx"

_SECT_XML = """<w:sectPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:pgSz w:w="11906" w:h="16838"/>
  <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"
           w:header="708" w:footer="708" w:gutter="0"/>
  <w:cols w:space="708"/>
  <w:docGrid w:linePitch="360"/>
</w:sectPr>"""


def ensure_terminal_sect_pr(doc: Document) -> None:
    """Some packaged DOCX omit body-level ``w:sectPr``; docxcompose then crashes."""
    body = doc.element.body
    for ch in body.iterchildren():
        if ch.tag == qn("w:sectPr"):
            return
    body.append(parse_xml(_SECT_XML))


def finalize_fragment_for_compose(path: Path) -> None:
    d = Document(str(path))
    ensure_terminal_sect_pr(d)
    d.save(str(path))


def _load_build_module():
    p = ROOT / "scripts" / "build_final_sdd_docx.py"
    spec = importlib.util.spec_from_file_location("bfd", p)
    if spec is None or spec.loader is None:
        sys.exit(f"cannot load {p}")
    m = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(m)
    return m


def extract_section32_md(md_text: str) -> str:
    m_start = re.search(r"(?ms)^## 3\.2 Files\s*$", md_text)
    m_end = re.search(r"(?ms)^## 3\.3 ", md_text)
    if not m_start or not m_end or m_end.start() <= m_start.start():
        sys.exit("ipcs_sdd.md: could not isolate ## 3.2 Files .. ## 3.3 boundary")
    return md_text[m_start.start() : m_end.start()].rstrip() + "\n"


def heading_context(para: Paragraph) -> tuple[str, str]:
    nm = para.style.name if para.style else ""
    return nm, (para.text or "").strip()


def find_section32_bounds(doc: Document):
    """``(first 3.2 body block, Heading 2 '3.3 … External')``.

    ``3.2 Files`` may be *Normal* (not Heading 2); TOC uses *Compact*.
    After Heading 2 ``3.1 … Definition``, the first ``w:tbl`` or ``w:p`` whose
    text begins with ``3.2`` starts the replace span.
    """

    elems = list(doc.element.body.iterchildren())
    anchor_el = None
    i31: int | None = None

    for idx, el in enumerate(elems):
        if el.tag != qn("w:p"):
            continue
        p = Paragraph(el, doc)
        st, txt = heading_context(p)
        if i31 is None and st == "Heading 2" and "3.1" in txt and "Definition" in txt:
            i31 = idx
            continue
        if (
            i31 is not None
            and anchor_el is None
            and st == "Heading 2"
            and txt.startswith("3.3")
            and "External" in txt
        ):
            anchor_el = el
            break

    if i31 is None or anchor_el is None:
        return None, None

    anchor_idx = elems.index(anchor_el)
    start_el = None
    for el in elems[i31 + 1 : anchor_idx]:
        if el.tag == qn("w:tbl"):
            start_el = el
            break
        if el.tag == qn("w:p"):
            p = Paragraph(el, doc)
            _, txt = heading_context(p)
            if txt.startswith("3.2"):
                start_el = el
                break

    return start_el, anchor_el


def section32_cut_indices(doc: Document) -> tuple[int, int] | None:
    """Return ``(start_idx, anchor_idx)`` into ``body`` children, or None."""
    start_el, anchor_el = find_section32_bounds(doc)
    if start_el is None or anchor_el is None:
        return None
    elems = list(doc.element.body.iterchildren())
    si = elems.index(start_el)
    ai = elems.index(anchor_el)
    if ai <= si:
        sys.exit("[error] invalid cut: 3.3 anchor precedes 3.2 start")
    return si, ai


def save_trimmed_head(master_path: Path, start_idx: int, out_path: Path) -> None:
    doc = Document(str(master_path))
    body = doc.element.body
    elems = list(body.iterchildren())
    for el in elems[start_idx:]:
        body.remove(el)
    doc.save(str(out_path))


def save_trimmed_tail(master_path: Path, anchor_idx: int, out_path: Path) -> None:
    doc = Document(str(master_path))
    body = doc.element.body
    elems = list(body.iterchildren())
    for el in elems[:anchor_idx]:
        body.remove(el)
    doc.save(str(out_path))


def run_patch() -> None:
    bfd = _load_build_module()

    if not FINAL_DOCX.is_file():
        sys.exit(f"missing {FINAL_DOCX}")
    if not SOURCE_MD.is_file():
        sys.exit(f"missing {SOURCE_MD}")

    shutil.copyfile(FINAL_DOCX, WORK_COPY)

    master = Document(str(WORK_COPY))
    ensure_terminal_sect_pr(master)
    master.save(str(WORK_COPY))

    probe = Document(str(WORK_COPY))
    cut = section32_cut_indices(probe)
    if cut is None:
        bfd._unlink_retry(WORK_COPY)
        sys.exit(
            "[error] Cannot locate 3.2 body span (first 3.2 content after '3.1 Definition' "
            "through before Heading 2 '3.3 … External'); check final_sdd.docx structure."
        )
    si, ai = cut

    frag_raw = SOURCE_MD.read_text(encoding="utf-8")
    frag_md = extract_section32_md(frag_raw)
    frag_md = bfd.rasterize_svg_refs_for_docx(frag_md, clear_cache=False)
    TMP_MD.write_text(frag_md, encoding="utf-8")

    bfd.run_pandoc(TMP_MD, TMP_DOCX)

    save_trimmed_head(WORK_COPY, si, PATCH_HEAD)
    save_trimmed_tail(WORK_COPY, ai, PATCH_TAIL)

    finalize_fragment_for_compose(PATCH_HEAD)
    finalize_fragment_for_compose(PATCH_TAIL)

    comp = Composer(Document(str(PATCH_HEAD)))
    comp.append(Document(str(TMP_DOCX)))
    comp.append(Document(str(PATCH_TAIL)))

    out_tmp = Path(tempfile.mkstemp(suffix=".docx", dir=str(ROOT))[1])
    try:
        comp.save(str(out_tmp))
        shutil.copyfile(str(out_tmp), str(FINAL_DOCX))
    finally:
        bfd._unlink_retry(out_tmp)
        bfd._unlink_retry(TMP_MD)
        bfd._unlink_retry(TMP_DOCX)
        bfd._unlink_retry(PATCH_HEAD)
        bfd._unlink_retry(PATCH_TAIL)
        bfd._unlink_retry(WORK_COPY)

    print(f"[done] patched 3.2 Files span in → {FINAL_DOCX}")


if __name__ == "__main__":
    run_patch()
