# -*- coding: utf-8 -*-
"""Post-process files_32_svgs and refresh 3.2 images in final_sdd.docx (SVG parts, not PNG)."""
from __future__ import annotations

import re
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml.shape import CT_Inline
from docx.parts.image import ImagePart
from docx.shared import Cm, Emu, Length

WORK = Path(__file__).resolve().parent
SVG_DIR = WORK / "files_32_svgs"
DOCX = WORK / "final_sdd.docx"
DOCX_FALLBACK = WORK / "final_sdd_patched.docx"

NUM_TO_SVG = {
    2: "3_2_2",
    3: "3_2_3",
    4: "3_2_4",
    5: "3_2_5",
    6: "3_2_6",
    7: "3_2_7",
    8: "3_2_8",
    9: "3_2_9",
    10: "3_2_10",
    11: "3_2_11",
    13: "3_2_13",
    14: "3_2_14",
    15: "3_2_15",
    16: "3_2_16",
    17: "3_2_17",
}


def clean_component_svg(raw: bytes) -> bytes:
    text = raw.decode("utf-8", errors="replace")
    text = re.sub(r"<\?plantuml[^?]*\?>", "", text)
    text = re.sub(r"<!--SRC=\[[^\]]*\]-->", "", text)
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    root = BeautifulSoup(text, "xml")
    svg = root.find("svg")
    if not svg:
        return raw
    for ttl in svg.find_all("title"):
        ttl.decompose()
    for g in svg.find_all("g"):
        cls = g.get("class") or []
        if isinstance(cls, str):
            cls = cls.split()
        gid = g.get("id") or ""
        is_link = (
            "link" in cls
            or gid.startswith("lnk")
            or gid.startswith("link_")
            or g.get("data-link-type")
        )
        if not is_link:
            continue
        for ch in list(g.children):
            if not getattr(ch, "name", None):
                continue
            nm = ch.name.lower()
            if nm == "text":
                ch.extract()
            elif nm == "polygon" and ch.get("fill") == "none":
                ch.extract()
            elif nm == "g" and ch.get("data-visibility-modifier"):
                ch.extract()
    return str(svg).encode("utf-8")


def refresh_svgs_on_disk():
    for p in sorted(SVG_DIR.glob("3_2_*.svg")):
        cleaned = clean_component_svg(p.read_bytes())
        p.write_bytes(cleaned)


def svg_dimensions(svg_bytes: bytes) -> tuple[int, int]:
    t = svg_bytes.decode("utf-8", errors="replace")
    m = re.search(r"viewBox\s*=\s*[\"']([^\"']+)[\"']", t, re.I)
    if m:
        parts = re.split(r"[\s,]+", m.group(1).strip())
        if len(parts) >= 4:
            w, h = float(parts[2]), float(parts[3])
            return max(1, int(round(w))), max(1, int(round(h)))
    m = re.search(r'\bwidth\s*=\s*["\'](\d+(?:\.\d+)?)px', t, re.I)
    w = float(m.group(1)) if m else None
    m = re.search(r'\bheight\s*=\s*["\'](\d+(?:\.\d+)?)px', t, re.I)
    h = float(m.group(1)) if m else None
    if w and h:
        return int(round(w)), int(round(h))
    return 400, 300


def add_svg_to_run(run, svg_bytes: bytes, filename: str, width: Length | None) -> None:
    part = run.part
    package = part.package
    px_w, px_h = svg_dimensions(svg_bytes)
    emu_per_px = 914400 / 96.0
    cx = int(round(px_w * emu_per_px))
    cy = int(round(px_h * emu_per_px))
    if width is not None:
        w_int = int(width)
        scale = w_int / cx
        cx = w_int
        cy = int(round(cy * scale))
    partname = package.image_parts._next_image_partname("svg")
    image_part = ImagePart(partname, "image/svg+xml", svg_bytes)
    package.image_parts.append(image_part)
    r_id = part.relate_to(image_part, RT.IMAGE)
    shape_id = part.next_id
    inline = CT_Inline.new_pic_inline(shape_id, r_id, filename, Emu(cx), Emu(cy))
    run._r.add_drawing(inline)


def iter_body_paragraphs(doc: Document):
    for child in doc.element.body:
        if child.tag.endswith("p"):
            yield child


def paragraph_text(p_el) -> str:
    return "".join(t.text or "" for t in p_el.iter(qn("w:t"))).strip()


def paragraph_has_drawing(p_el) -> bool:
    return bool(p_el.findall(".//" + qn("w:drawing")))


def replace_paragraph_with_svg(
    doc: Document,
    p_el,
    svg_bytes: bytes,
    filename: str,
    width_cm: float = 15.0,
) -> None:
    parent = p_el.getparent()
    idx = list(parent).index(p_el)
    parent.remove(p_el)
    np = doc.add_paragraph()
    run = np.add_run()
    add_svg_to_run(run, svg_bytes, filename, Cm(width_cm))
    new_el = np._element
    new_el.getparent().remove(new_el)
    parent.insert(idx, new_el)


def patch_section_32_images(doc: Document):
    in_files = False
    pending: tuple[bytes, str] | None = None
    for p_el in list(iter_body_paragraphs(doc)):
        t = paragraph_text(p_el)
        if re.match(r"^3\.2\.1\s", t):
            in_files = True
            pending = None
            continue
        if in_files and t.startswith("3.3 External"):
            break
        m = re.match(r"^3\.2\.(\d+)\s", t)
        if in_files and m:
            n = int(m.group(1))
            if n in NUM_TO_SVG:
                svg_path = SVG_DIR / f"{NUM_TO_SVG[n]}.svg"
                if svg_path.is_file():
                    blob = clean_component_svg(svg_path.read_bytes())
                    pending = (blob, svg_path.name)
                else:
                    pending = None
            else:
                pending = None
            continue
        if in_files and pending and paragraph_has_drawing(p_el):
            b, fn = pending
            replace_paragraph_with_svg(doc, p_el, b, fn)
            pending = None


def main():
    refresh_svgs_on_disk()
    src = DOCX if DOCX.is_file() else DOCX_FALLBACK
    if not src.is_file():
        raise SystemExit("No final_sdd.docx or final_sdd_patched.docx found")
    doc = Document(str(src))
    patch_section_32_images(doc)
    try:
        doc.save(str(DOCX))
    except PermissionError:
        doc.save(str(DOCX_FALLBACK))
        print(
            f"NOTE: {DOCX.name} locked; saved {DOCX_FALLBACK.name}. "
            "Close Word and replace manually."
        )


if __name__ == "__main__":
    main()
