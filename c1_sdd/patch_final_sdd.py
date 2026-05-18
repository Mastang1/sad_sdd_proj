# -*- coding: utf-8 -*-
"""Patch final_sdd.docx: 3.2 UML figures + 3.3/3.4 function tables from ipcs_sdd.md."""

from __future__ import annotations

import io
import re
from pathlib import Path

import cairosvg
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table

WORK = Path(__file__).resolve().parent
DOCX_IN = WORK / "final_sdd.docx"
DOCX_OUT = WORK / "final_sdd_patched.docx"
MD_SRC = WORK / "ipcs_sdd.md"
SVG_DIR = WORK / "files_32_svgs"

# 3_subsections that have component UML SVGs (no 3.2.12 in source set)
NUM_TO_SVG = {
    2: "3_2_2",
    3: "3_2_3",
    4: "3_2_4",
    5: "3_2_5",
    6: "3_2_6",
    7: "3_2_7",
    8: "3_2_8",
    9: "3_2_9",
    10: "3_2_10",
    11: "3_2_11",
    13: "3_2_13",
    14: "3_2_14",
    15: "3_2_15",
    16: "3_2_16",
    17: "3_2_17",
}

SVG_NS = "http://www.w3.org/2000/svg"


def clean_svg_bytes(raw: bytes) -> bytes:
    """Remove edge labels, hollow diamonds on links, PlantUML junk; normalize khaki tones."""
    text = raw.decode("utf-8", errors="replace")
    text = re.sub(r"<!--SRC=\[[^\]]*\]-->", "", text)
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    text = text.replace("#F5E6CC", "#F0E68C").replace("#FFFCF5", "#FFF8DC")
    root = BeautifulSoup(text, "xml")
    svg = root.find("svg")
    if not svg:
        return raw
    # Title: keep simple ASCII to avoid Word/font issues
    ttl = svg.find("title")
    if ttl:
        nt = re.sub(r"\s+", " ", ttl.get_text() or "")[:120]
        ttl.clear()
        ttl.append(nt)
    for g in svg.find_all("g"):
        gid = g.get("id") or ""
        if not gid.startswith("link_"):
            continue
        for ch in list(g.children):
            if not getattr(ch, "name", None):
                continue
            nm = ch.name.lower()
            if nm == "text":
                ch.extract()
            elif nm == "polygon" and ch.get("fill") == "none":
                ch.extract()
    return str(svg).encode("utf-8")


def svg_to_png(svg_bytes: bytes) -> bytes:
    return cairosvg.svg2png(bytestring=svg_bytes, dpi=120)


def iter_body_paragraphs(doc: Document):
    for child in doc.element.body:
        if child.tag.endswith("p"):
            yield child


def paragraph_text(p_el) -> str:
    return "".join(t.text or "" for t in p_el.iter(qn("w:t"))).strip()


def paragraph_has_drawing(p_el) -> bool:
    return bool(p_el.findall(".//" + qn("w:drawing")))


def replace_paragraph_with_image(doc: Document, p_el, png_bytes: bytes, width_cm: float = 15.0):
    parent = p_el.getparent()
    idx = list(parent).index(p_el)
    parent.remove(p_el)
    np = doc.add_paragraph()
    nr = np.add_run()
    stream = io.BytesIO(png_bytes)
    nr.add_picture(stream, width=Cm(width_cm))
    new_el = np._element
    new_el.getparent().remove(new_el)
    parent.insert(idx, new_el)


def load_md_tables() -> dict[str, str]:
    md = MD_SRC.read_text(encoding="utf-8")
    pat = r"### (3\.[34]\.\d+ \w+)\s*\n+(<table.*?</table>)"
    found = re.findall(pat, md, flags=re.DOTALL)
    return dict(found)


def parse_html_placements(table_html: str):
    soup = BeautifulSoup(table_html, "html.parser")
    table = soup.find("table")
    if not table:
        return 0, 0, []
    occupied: dict[tuple[int, int], bool] = {}
    placements: list[tuple[int, int, int, int, str]] = []
    cur_r = 0
    for tr in table.find_all("tr"):
        cur_c = 0
        tds = tr.find_all("td")
        if not tds:
            continue
        for td in tds:
            while (cur_r, cur_c) in occupied:
                cur_c += 1
            rs = int(td.get("rowspan", 1))
            cs = int(td.get("colspan", 1))
            text = td.get_text(separator="\n").strip()
            placements.append((cur_r, cur_c, rs, cs, text))
            for dr in range(rs):
                for dc in range(cs):
                    occupied[(cur_r + dr, cur_c + dc)] = True
            cur_c += cs
        cur_r += 1
    nrows = cur_r
    ncols = max((c + cs for _, c, _, cs, _ in placements), default=0)
    return nrows, ncols, placements


def insert_table_after_element(doc: Document, anchor_p_el, nrows: int, ncols: int, placements):
    temp = doc.add_table(rows=nrows, cols=ncols)
    tbl_el = temp._tbl
    temp._tbl.getparent().remove(tbl_el)
    anchor_p_el.addnext(tbl_el)
    table = Table(tbl_el, doc)

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    for r, c, rs, cs, text in placements:
        cell = table.cell(r, c)
        if rs > 1 or cs > 1:
            cell.merge(table.cell(r + rs - 1, c + cs - 1))
        cell.text = text or " "
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
    return table


def remove_elements_between(body, start_el, end_el):
    """Remove contiguous siblings after start_el up to but not including end_el."""
    parent = start_el.getparent()
    kids = list(parent)
    i0 = kids.index(start_el) + 1
    i1 = kids.index(end_el)
    for el in kids[i0:i1]:
        parent.remove(el)


def has_table_between(body, start_el, end_el):
    kids = list(body)
    try:
        i0 = kids.index(start_el) + 1
        i1 = kids.index(end_el)
    except ValueError:
        return False
    for el in kids[i0:i1]:
        if el.tag.split("}")[-1] == "tbl":
            return True
    return False


def patch_function_tables(doc: Document, md_tables: dict[str, str]):
    body = doc.element.body
    children = list(body)
    jobs = []
    for i, ch in enumerate(children):
        if ch.tag.split("}")[-1] != "p":
            continue
        title = paragraph_text(ch)
        if not (title.startswith("3.3.") or title.startswith("3.4.")):
            continue
        parts = title.split()
        if len(parts) < 2:
            continue
        if "processing" in parts[1] or "flow" in title:
            continue
        key = f"{parts[0]} {parts[1]}"
        if key not in md_tables:
            continue
        end_el = None
        for j in range(i + 1, len(children)):
            cj = children[j]
            if cj.tag.split("}")[-1] != "p":
                continue
            tj = paragraph_text(cj)
            if re.match(r"3\.3\.\d+ processing flow$", tj) or re.match(
                r"3\.4\.\d+ processing flow$", tj
            ):
                end_el = cj
                break
            if not tj and paragraph_has_drawing(cj):
                end_el = cj
                break
        if end_el is None:
            continue
        if has_table_between(body, ch, end_el):
            continue
        jobs.append((ch, md_tables[key], end_el))

    for ch, table_html, end_el in reversed(jobs):
        nrows, ncols, placements = parse_html_placements(table_html)
        if nrows == 0 or ncols == 0:
            continue
        remove_elements_between(body, ch, end_el)
        insert_table_after_element(doc, ch, nrows, ncols, placements)


def patch_section_32_images(doc: Document):
    in_files = False
    pending_png: bytes | None = None
    for p_el in list(iter_body_paragraphs(doc)):
        t = paragraph_text(p_el)
        # Skip TOC: only enter Files body after first 3.2.1 (file list) subsection
        if re.match(r"^3\.2\.1\s", t):
            in_files = True
            pending_png = None
            continue
        if in_files and t.startswith("3.3 External"):
            break
        m = re.match(r"^3\.2\.(\d+)\s", t)
        if in_files and m:
            n = int(m.group(1))
            if n in NUM_TO_SVG:
                svg_path = SVG_DIR / f"{NUM_TO_SVG[n]}.svg"
                if svg_path.is_file():
                    raw = svg_path.read_bytes()
                    pending_png = svg_to_png(clean_svg_bytes(raw))
                else:
                    pending_png = None
            else:
                pending_png = None
            continue
        if in_files and pending_png and paragraph_has_drawing(p_el):
            replace_paragraph_with_image(doc, p_el, pending_png)
            pending_png = None


def main():
    import sys

    src = Path(sys.argv[1]) if len(sys.argv) > 1 else DOCX_IN

    md_tables = load_md_tables()
    if len(md_tables) != 82:
        raise SystemExit(f"expected 82 md tables, got {len(md_tables)}")

    doc = Document(str(src))
    patch_section_32_images(doc)
    patch_function_tables(doc, md_tables)
    try:
        doc.save(str(DOCX_IN))
    except PermissionError:
        dest = src if src.name != "final_sdd.docx" else DOCX_OUT
        doc.save(str(dest))
        print(
            f"NOTE: {DOCX_IN.name} is locked; saved to {dest.name}. "
            "Close Word and copy/rename to replace the original."
        )


if __name__ == "__main__":
    main()
